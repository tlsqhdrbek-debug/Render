import streamlit as st
import fitz  # PyMuPDF
import os
import re
import numpy as np
from openai import OpenAI
from pathlib import Path
from docx import Document
from datetime import datetime
import tiktoken
import json
import traceback
import time
import uuid
from functools import wraps
import pandas as pd

# 페이지 설정
st.set_page_config(
    page_title="기업 분석 보고서 생성기",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 커스텀 CSS
st.markdown("""
<style>
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
    }
    .stApp {
        background: white;
        border-radius: 24px;
        padding: 40px;
        margin: 20px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.15);
    }
    /* 사이드바 너비 증가 */
    [data-testid="stSidebar"] {
        min-width: 350px !important;
        max-width: 350px !important;
    }
    /* 사이드바 버튼 크기 조정 */
    .stSidebar button {
        font-size: 11px !important;
        padding: 5px 10px !important;
        white-space: nowrap !important;
    }
    /* Expander 내부 버튼도 작게 */
    [data-testid="stExpander"] button {
        font-size: 11px !important;
    }
    /* 탭 스타일 개선 */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: #f8fafc;
        padding: 10px;
        border-radius: 12px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        padding: 0 30px;
        background-color: white;
        border-radius: 8px;
        font-size: 16px !important;
        font-weight: 600 !important;
        color: #64748b;
        border: 2px solid transparent;
        transition: all 0.3s ease;
    }
    .stTabs [data-baseweb="tab"]:hover {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.3);
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        border-color: transparent !important;
    }
    .keyword-tag {
        display: inline-flex;
        align-items: center;
        padding: 4px 10px;
        margin: 3px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 12px;
        font-size: 12px;
        font-weight: 500;
        white-space: nowrap;
    }
    .template-container {
        display: flex;
        flex-wrap: wrap;
        gap: 5px;
        align-items: center;
        margin-bottom: 10px;
    }
    .delete-btn {
        background: rgba(255, 255, 255, 0.2);
        border: none;
        color: white;
        width: 20px;
        height: 20px;
        border-radius: 50%;
        cursor: pointer;
        margin-left: 8px;
    }
</style>
""", unsafe_allow_html=True)

# .env 파일 로드
def load_env():
    """Load environment variables from .env file"""
    env_path = Path(__file__).parent / ".env"
    if env_path.exists():
        with open(env_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    os.environ[key.strip()] = value.strip()

load_env()

# OpenAI 클라이언트 초기화
openai_client = None
if os.getenv("OPENAI_API_KEY") and os.getenv("OPENAI_API_KEY") != "your-api-key-here":
    openai_api_key = os.getenv("OPENAI_API_KEY")
    openai_client = OpenAI(api_key=openai_api_key)
    st.sidebar.success("✅ OpenAI API 키 로드됨")
else:
    st.sidebar.warning("⚠️ .env에 OPENAI_API_KEY를 설정하세요")

# Supabase 클라이언트 초기화
supabase_client = None
try:
    from supabase import create_client, Client
    supabase_url = os.getenv("SUPABASE_URL")
    supabase_key = os.getenv("SUPABASE_KEY")
    
    if supabase_url and supabase_key:
        supabase_client = create_client(supabase_url, supabase_key)
        st.sidebar.success("✅ Supabase 연결됨")
    else:
        st.sidebar.info("ℹ️ Supabase 미연결 (환경변수 미설정)")
except Exception as e:
    st.sidebar.error(f"⚠️ Supabase 연결 실패: {e}")

# Upstage API 키 확인
UPSTAGE_API_KEY = os.getenv("UPSTAGE_API_KEY")
if UPSTAGE_API_KEY and UPSTAGE_API_KEY != "your-upstage-api-key-here":
    st.sidebar.success("✅ Upstage Document Parse 연결됨")
else:
    st.sidebar.info("ℹ️ Upstage API 미설정 (기본 텍스트 추출)")

# 세션 스테이트 초기화
if 'template' not in st.session_state:
    st.session_state.template = []
if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = {}
if 'pdf_text' not in st.session_state:
    st.session_state.pdf_text = ""
if 'report_sections' not in st.session_state:
    # 기본 보고서 섹션 선택 (모두 선택)
    st.session_state.report_sections = [
        "기업 개요",
        "사업 구조 및 Revenue Model 분석",
        "산업 및 시장 분석",
        "재무 요약",
        "재무 건전성 심화 분석",
        "고객사 및 매출 집중도 분석",
        "경쟁사 비교 분석",
        "경영진 역량 및 지배구조 분석",
        "신용도 및 법률 리스크",
        "리스크 요인",
        "종합 평가"
    ]
if 'show_template_editor' not in st.session_state:
    st.session_state.show_template_editor = False
if 'reference_pdfs' not in st.session_state:
    st.session_state.reference_pdfs = {}  # {filename: extracted_text}
if 'structured_data' not in st.session_state:
    st.session_state.structured_data = None  # Upstage Parse 구조화 데이터

# 로깅 시스템용 세션 스테이트
if 'user_name' not in st.session_state:
    st.session_state.user_name = None
if 'user_email' not in st.session_state:
    st.session_state.user_email = None
if 'session_id' not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if 'current_test_session_id' not in st.session_state:
    st.session_state.current_test_session_id = None

# ============================================
# 로깅 시스템
# ============================================

def create_or_get_test_user(name, email=None):
    """테스트 사용자 생성 또는 조회"""
    if not supabase_client:
        return None
    
    try:
        # 이름과 이메일로 기존 사용자 확인
        if email:
            response = supabase_client.table("test_users").select("*").eq("name", name).eq("email", email).execute()
        else:
            response = supabase_client.table("test_users").select("*").eq("name", name).is_("email", "null").execute()
        
        if response.data:
            return response.data[0]
        
        # 새 사용자 생성
        user_data = {
            "name": name,
            "email": email
        }
        response = supabase_client.table("test_users").insert(user_data).execute()
        return response.data[0] if response.data else None
    except Exception as e:
        st.error(f"사용자 생성 실패: {e}")
        return None

def start_test_session(user_id, company_name, pdf_filename):
    """테스트 세션 시작"""
    if not supabase_client:
        return None
    
    try:
        session_data = {
            "user_id": user_id,
            "company_name": company_name,
            "pdf_filename": pdf_filename,
            "status": "in_progress"
        }
        response = supabase_client.table("test_sessions").insert(session_data).execute()
        session_id = response.data[0]["id"] if response.data else None
        st.session_state.current_test_session_id = session_id
        return session_id
    except Exception as e:
        st.error(f"세션 시작 실패: {e}")
        return None

def complete_test_session(status, error_message=None, execution_time_ms=None):
    """테스트 세션 완료"""
    if not supabase_client or not st.session_state.current_test_session_id:
        return
    
    try:
        update_data = {
            "completed_at": datetime.now().isoformat(),
            "status": status
        }
        if error_message:
            update_data["error_message"] = error_message
        # execution_time_ms는 저장하지 않음 (테이블에 필드 없음)
        
        supabase_client.table("test_sessions").update(update_data).eq("id", st.session_state.current_test_session_id).execute()
    except Exception as e:
        st.error(f"세션 완료 기록 실패: {e}")

def log_activity(step, status, details=None, execution_time_ms=None):
    """활동 로그 기록"""
    if not supabase_client:
        return
    
    # user_login은 세션 ID 없이도 기록 (일반 텍스트 session_id 사용)
    if step == "user_login":
        try:
            log_data = {
                "session_id": st.session_state.session_id,  # UUID가 아닌 일반 텍스트 세션 ID
                "step": step,
                "status": status,
                "details": details if details else {},
                "execution_time_ms": execution_time_ms
            }
            supabase_client.table("activity_logs").insert(log_data).execute()
        except Exception as e:
            print(f"로그 기록 실패: {e}")
        return
    
    # 그 외 로그는 test_session_id 필요
    if not st.session_state.current_test_session_id:
        return
    
    try:
        log_data = {
            "session_id": st.session_state.current_test_session_id,
            "step": step,
            "status": status,
            "details": details if details else {},
            "execution_time_ms": execution_time_ms
        }
        supabase_client.table("activity_logs").insert(log_data).execute()
    except Exception as e:
        print(f"로그 기록 실패: {e}")  # st.error 대신 print 사용 (로그 기록 실패는 사용자에게 노출 안 함)

def log_error(step, error, stack_trace=None):
    """에러 로그 기록"""
    if not supabase_client:
        return
    
    try:
        error_details = {
            "error_type": type(error).__name__,
            "error_message": str(error),
            "stack_trace": stack_trace or traceback.format_exc()
        }
        
        # test_session_id가 있으면 UUID 사용, 없으면 일반 session_id 사용
        session_id_to_use = st.session_state.current_test_session_id if st.session_state.current_test_session_id else st.session_state.session_id
        
        log_data = {
            "session_id": session_id_to_use,
            "step": step,
            "status": "failed",
            "details": error_details
        }
        supabase_client.table("activity_logs").insert(log_data).execute()
    except Exception as e:
        print(f"에러 로그 기록 실패: {e}")

def log_execution_time(step_name):
    """실행 시간 측정 데코레이터"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            start_time = time.time()
            log_activity(step_name, "started")
            
            try:
                result = func(*args, **kwargs)
                execution_time_ms = int((time.time() - start_time) * 1000)
                log_activity(step_name, "success", execution_time_ms=execution_time_ms)
                return result
            except Exception as e:
                execution_time_ms = int((time.time() - start_time) * 1000)
                log_error(step_name, e)
                log_activity(step_name, "failed", execution_time_ms=execution_time_ms)
                raise
        return wrapper
    return decorator

def log_data_quality(
    selected_keywords,
    ocr_raw_text,
    ocr_structured_data,
    llm_extracted_data,
    llm_extraction_time_ms,
    company_name=None,
    pdf_filename=None,
    pdf_pages=None,
    report_content=None,
    report_generation_time_ms=None
):
    """데이터 품질 검증 로그 기록 - OCR, LLM 추출, 보고서 생성 비교"""
    if not supabase_client or not st.session_state.current_test_session_id:
        return None
    
    try:
        # 추출 성공률 계산
        keywords_with_data = sum(1 for v in llm_extracted_data.values() if v and v != "정보 없음")
        keywords_missing_data = len(llm_extracted_data) - keywords_with_data
        extraction_success_rate = (keywords_with_data / len(llm_extracted_data) * 100) if llm_extracted_data else 0
        
        # 표 및 차트 개수 계산
        ocr_tables_count = len(ocr_structured_data.get('tables', [])) if ocr_structured_data else 0
        ocr_charts_count = len(ocr_structured_data.get('charts', [])) if ocr_structured_data else 0
        
        log_data = {
            "session_id": st.session_state.current_test_session_id,
            "user_name": st.session_state.user_name,
            "company_name": company_name,
            
            # 1. 선택된 키워드
            "selected_keywords": selected_keywords,
            
            # 2. OCR 원본 데이터
            "ocr_raw_text": ocr_raw_text[:20000] if ocr_raw_text else None,  # 처음 20000자만 저장 (재무표 전체 포함)
            "ocr_structured_data": ocr_structured_data,
            "ocr_tables_count": ocr_tables_count,
            "ocr_charts_count": ocr_charts_count,
            
            # 3. LLM 추출 데이터
            "llm_extracted_data": llm_extracted_data,
            "llm_model": "gpt-4o-mini",
            "llm_extraction_time_ms": llm_extraction_time_ms,
            
            # 4. 보고서 데이터
            "report_generated": report_content is not None,
            "report_content": report_content[:20000] if report_content else None,  # 처음 20000자만 저장
            "report_model": "gpt-4o-mini" if report_content else None,
            "report_generation_time_ms": report_generation_time_ms,
            
            # 5. 품질 메트릭
            "extraction_success_rate": round(extraction_success_rate, 2),
            "keywords_with_data": keywords_with_data,
            "keywords_missing_data": keywords_missing_data,
            
            # 6. 기타
            "pdf_filename": pdf_filename,
            "pdf_pages": pdf_pages,
        }
        
        result = supabase_client.table("data_quality_logs").insert(log_data).execute()
        
        if result.data and len(result.data) > 0:
            log_id = result.data[0]['id']
            print(f"✅ 데이터 품질 로그 저장 완료: {log_id}")
            return log_id
        
        return None
        
    except Exception as e:
        print(f"❌ 데이터 품질 로그 저장 실패: {e}")
        traceback.print_exc()
        return None

def generate_quality_log_txt(log_data):
    """데이터 품질 로그를 AI 분석용 TXT 파일로 변환"""
    
    # 기본 정보
    company = log_data.get('company_name', 'Unknown')
    user = log_data.get('user_name', 'N/A')
    created_at = log_data.get('created_at', 'N/A')[:19].replace('T', ' ')
    pdf_filename = log_data.get('pdf_filename', 'N/A')
    pdf_pages = log_data.get('pdf_pages', 0)
    
    # 품질 메트릭
    success_rate = log_data.get('extraction_success_rate', 0)
    keywords_success = log_data.get('keywords_with_data', 0)
    keywords_failed = log_data.get('keywords_missing_data', 0)
    ocr_tables = log_data.get('ocr_tables_count', 0)
    ocr_charts = log_data.get('ocr_charts_count', 0)
    
    # 키워드, OCR, LLM 데이터
    keywords = log_data.get('selected_keywords', [])
    ocr_raw = log_data.get('ocr_raw_text', '')
    structured_data = log_data.get('ocr_structured_data', {})
    extracted = log_data.get('llm_extracted_data', {})
    report = log_data.get('report_content', '')
    
    # TXT 파일 생성
    txt = []
    txt.append("=" * 80)
    txt.append("데이터 품질 검증 로그 - AI 분석용")
    txt.append("=" * 80)
    txt.append("")
    txt.append("[기본 정보]")
    txt.append(f"- 회사명: {company}")
    txt.append(f"- 사용자: {user}")
    txt.append(f"- 작성일: {created_at}")
    txt.append(f"- PDF 파일: {pdf_filename}")
    txt.append(f"- PDF 페이지: {pdf_pages}페이지")
    txt.append("")
    txt.append("[품질 메트릭]")
    txt.append(f"- 추출 성공률: {success_rate}%")
    txt.append(f"- 성공: {keywords_success}개")
    txt.append(f"- 실패: {keywords_failed}개")
    txt.append(f"- OCR 표 인식: {ocr_tables}개")
    txt.append(f"- OCR 차트/그래프 인식: {ocr_charts}개")
    txt.append("")
    
    # 1. 선택된 키워드
    txt.append("=" * 80)
    txt.append("1. 선택된 추출 키워드")
    txt.append("=" * 80)
    txt.append("")
    if keywords:
        for idx, kw in enumerate(keywords, 1):
            txt.append(f"{idx}. {kw}")
        txt.append("")
        txt.append(f"(총 {len(keywords)}개 키워드)")
    else:
        txt.append("키워드 정보 없음")
    txt.append("")
    txt.append("")
    
    # 2. OCR 원본 데이터
    txt.append("=" * 80)
    txt.append("2. OCR 원본 데이터 (Upstage Parse)")
    txt.append("=" * 80)
    txt.append("")
    
    # 표 데이터
    if structured_data and structured_data.get('tables'):
        txt.append(f"[표 데이터 - 총 {len(structured_data['tables'])}개]")
        txt.append("")
        for idx, table in enumerate(structured_data['tables'], 1):  # 모든 표 표시
            txt.append(f"--- 표 {idx} (페이지 {table.get('page', '?')}) ---")
            table_content = table.get('content', '내용 없음')
            txt.append(table_content[:1000])  # 각 표당 1000자로 증가 (재무표 전체 포함)
            if len(table_content) > 1000:
                txt.append("... (생략)")
            txt.append("")
    else:
        txt.append("[표 데이터 없음]")
        txt.append("")
    
    # 원본 텍스트
    txt.append("[추출된 원본 텍스트]")
    txt.append("")
    if ocr_raw:
        txt.append(ocr_raw[:5000])  # 처음 5000자로 증가 (재무표 전체 포함)
        if len(ocr_raw) > 5000:
            txt.append("")
            txt.append("... (이하 생략)")
    else:
        txt.append("원본 텍스트 없음")
    txt.append("")
    txt.append("")
    
    # 3. LLM 추출 데이터
    txt.append("=" * 80)
    txt.append("3. LLM 추출 데이터")
    txt.append("=" * 80)
    txt.append("")
    
    if extracted:
        # 성공/실패 구분
        success_data = {k: v for k, v in extracted.items() if v and v != "정보 없음"}
        failed_data = {k: v for k, v in extracted.items() if not v or v == "정보 없음"}
        
        txt.append(f"[✅ 성공적으로 추출된 데이터 - {len(success_data)}개]")
        txt.append("")
        if success_data:
            for idx, (key, value) in enumerate(success_data.items(), 1):
                txt.append(f"{idx}. {key}")
                txt.append(f"   → {value}")
                txt.append("")
        else:
            txt.append("없음")
            txt.append("")
        
        txt.append("")
        txt.append(f"[❌ 추출 실패 데이터 - {len(failed_data)}개]")
        txt.append("")
        if failed_data:
            for idx, key in enumerate(failed_data.keys(), 1):
                txt.append(f"{idx}. {key}")
                txt.append(f"   → 정보 없음")
                txt.append("")
        else:
            txt.append("없음")
            txt.append("")
    else:
        txt.append("LLM 추출 데이터 없음")
        txt.append("")
    
    txt.append("")
    txt.append(f"[LLM 처리 정보]")
    txt.append(f"- 모델: {log_data.get('llm_model', 'N/A')}")
    txt.append(f"- 처리 시간: {log_data.get('llm_extraction_time_ms', 0)}ms")
    txt.append("")
    txt.append("")
    
    # 4. 보고서 (선택)
    if log_data.get('report_generated') and report:
        txt.append("=" * 80)
        txt.append("4. 보고서 생성 결과 (선택)")
        txt.append("=" * 80)
        txt.append("")
        txt.append(report[:2000])  # 처음 2000자
        if len(report) > 2000:
            txt.append("")
            txt.append("... (이하 생략)")
        txt.append("")
        txt.append("")
        txt.append(f"[보고서 생성 정보]")
        txt.append(f"- 모델: {log_data.get('report_model', 'N/A')}")
        txt.append(f"- 생성 시간: {log_data.get('report_generation_time_ms', 0)}ms")
        txt.append(f"- 전체 길이: {len(report)}자")
        txt.append("")
        txt.append("")
    
    # AI 분석을 위한 질문
    txt.append("=" * 80)
    txt.append("AI 분석을 위한 질문")
    txt.append("=" * 80)
    txt.append("")
    txt.append("이 로그를 AI에게 첨부하고 다음과 같이 요청하세요:")
    txt.append("")
    txt.append("1. OCR 원본 데이터에는 있는데 LLM이 추출하지 못한 정보가 있나요?")
    txt.append("   → 어떤 키워드가 누락되었는지 구체적으로 분석해주세요.")
    txt.append("")
    txt.append("2. LLM이 잘못 추출한 값이 있나요?")
    txt.append("   → OCR 원본과 비교하여 잘못된 부분을 지적해주세요.")
    txt.append("   (예: 영업이익과 영업이익률 혼동, 단위 오류 등)")
    txt.append("")
    txt.append("3. 추출 실패 데이터에 대해:")
    txt.append("   → OCR 원본에 해당 정보가 있는지 확인해주세요.")
    txt.append("   → 있다면 왜 LLM이 찾지 못했는지 분석해주세요.")
    txt.append("")
    txt.append("4. 프롬프트를 어떻게 개선하면 추출 성공률을 높일 수 있나요?")
    txt.append("   → 구체적인 프롬프트 개선안을 제시해주세요.")
    txt.append("")
    txt.append("5. OCR 단계에서 개선이 필요한 부분이 있나요?")
    txt.append("   → 표 인식, 텍스트 추출 품질 등을 평가해주세요.")
    txt.append("")
    txt.append("")
    txt.append("=" * 80)
    txt.append("분석 완료 후 개선 방향")
    txt.append("=" * 80)
    txt.append("")
    txt.append("AI 분석 결과를 바탕으로:")
    txt.append("1. streamlit_app.py의 extract_all_keywords_batch() 함수 프롬프트 수정")
    txt.append("2. OCR 설정 조정 (표 구조 인식 모드 등)")
    txt.append("3. 키워드 정의 개선 (더 명확한 키워드명 사용)")
    txt.append("4. 재테스트 및 성공률 비교")
    txt.append("")
    txt.append("=" * 80)
    txt.append("파일 끝")
    txt.append("=" * 80)
    
    return "\n".join(txt)

# 보고서 섹션별 작성 지침 정의
REPORT_SECTION_TEMPLATES = {
    "기업 개요": """1. **기업 개요**
   - 회사명, 업종, 주요 사업 내용
   - 매출/영업이익 등 기본 정보
   - 설립 배경 및 주요 연혁
   - 조직 구조 요약(선택)
   - 추출된 데이터만 사용""",
    
    "사업 구조 및 Revenue Model 분석": """2. **사업 구조 및 Revenue Model 분석**
   - 주요 사업부 구조 및 매출 비중
   - 제품/서비스별 수익 모델
   - 고객 대상군(B2B/B2C), 지역별 매출 구조
   - 주요 원가/마진 구조
   - 여신 심사에 필요한 핵심 항목
   - 추출된 데이터 기반, 없으면 업종 특성 반영
   - *> 본 문서에 해당 내용이 없어, AI 모델이 학습한 일반 지식을 기반으로 작성했습니다. (2023년 10월 기준)*
   - *> 최신 정보나 경쟁사 비교가 필요하면 관련 PDF를 '참고자료'로 추가 업로드하시면 더 정확한 분석이 가능합니다.*""",
    
    "산업 및 시장 분석": """3. **산업 및 시장 분석**
   - 산업 규모, 성장성, 시장 동향
   - 경쟁 구도 및 트렌드
   - 해당 기업이 속한 시장의 위험 요인
   - PDF에 정보 없으면 업종 기준 일반 분석
   - *> 본 문서에 해당 내용이 없어, AI 모델이 학습한 일반 지식을 기반으로 작성했습니다. (2023년 10월 기준)*
   - *> 최신 산업 리포트나 시장 분석 자료를 '참고자료'로 추가 업로드하시면 더 정확한 분석이 가능합니다.*""",
    
    "재무 요약": """4. **재무 요약**
   - 매출, 영업이익, 성장률 분석
   - 추출된 핵심 재무 데이터 기반
   - 수익성·성장성 지표
   - 전년 대비 성장률, 수익성 지표 포함""",
    
    "재무 건전성 심화 분석": """5. **재무 건전성 심화 분석**
   - 부채 구조(단기/장기)
   - 이자보상배율, 차입 의존도
   - 영업현금흐름 안정성
   - 순운전자본(NWC) 분석
   - 대출 상환능력 평가 핵심 지표
   - 추출된 재무 데이터 기반""",
    
    "고객사 및 매출 집중도 분석": """6. **고객사 및 매출 집중도 분석**
   - 주요 고객사 TOP5
   - 단일 거래처 의존도
   - 매출 다변화 수준
   - 거래처 변경 가능성
   - 캐피탈 리스크 심사에서 매우 중요
   - PDF에 데이터가 있으면 사용, 없으면 언급""",
    
    "경쟁사 비교 분석": """7. **경쟁사 비교 분석**
   - 업종 내 주요 경쟁사 비교
   - 경쟁 우위·열위 분석
   - 시장 점유율 추정
   - 데이터 없으면 업종 기반 일반 비교
   - *> 본 문서에 해당 내용이 없어, AI 모델이 학습한 일반 지식을 기반으로 작성했습니다. (2023년 10월 기준)*
   - *> 경쟁사 사업보고서나 IR 자료를 '참고자료'로 추가 업로드하시면 더 정확한 비교 분석이 가능합니다.*""",
    
    "경영진 역량 및 지배구조 분석": """8. **경영진 역량 및 지배구조 분석**
   - CEO 및 핵심 임원 경력
   - 지분 구조, 오너 리스크
   - 지배구조 투명성
   - 경영진 교체 이력
   - 특히 중소기업 심사에 매우 중요
   - PDF에 데이터가 있으면 사용, 없으면 언급 안 함""",
    
    "신용도 및 법률 리스크": """9. **신용도 및 법률 리스크**
   - 신용등급(있으면)
   - 감사 의견(적정/한정 등)
   - 최근 소송·분쟁·제재 여부
   - 공정위/금융위 제재 여부
   - 기본 리스크 요인과 구분되는 정량적 리스크
   - PDF에 데이터가 있으면 사용
   - *> 본 문서에 해당 내용이 없어, AI 모델이 학습한 일반 지식을 기반으로 작성했습니다. (2023년 10월 기준)*""",
    
    "리스크 요인": """10. **리스크 요인**
   - 산업 리스크
   - 운영 리스크
   - 재무적 일반 리스크
   - PDF에 리스크 정보가 있으면 사용
   - 없으면 해당 산업의 일반적 리스크 설명
   - *> 본 문서에 해당 내용이 없어, AI 모델이 학습한 일반 지식을 기반으로 작성했습니다. (2023년 10월 기준)*""",
    
    "종합 평가": """11. **종합 평가 (투자/대출 관점)**
   - 재무 안정성 평가
   - 상환 능력 평가
   - 성장 가능성 요약
   - 종합 의견 및 권장 조치
   - 대출 승인/조건/유의사항 제시 가능
   - 추출된 데이터 기반으로 객관적 판단"""
}

# Supabase 헬퍼 함수
def save_to_supabase(company_name, pdf_file, extracted_text, extracted_data, report_content=None, create_embeddings_flag=True):
    """Supabase에 데이터 및 임베딩 저장"""
    if not supabase_client:
        st.warning("⚠️ Supabase 클라이언트가 연결되지 않았습니다.")
        return None
    
    try:
        # 1. 기업 정보 저장
        company_data = {
            "company_name": company_name,
            "industry": extracted_data.get("업종") or extracted_data.get("산업분류") or "미분류"
        }
        company_response = supabase_client.table("companies").insert(company_data).execute()
        company_id = company_response.data[0]["id"]
        st.info(f"✅ 기업 정보 저장 완료 (ID: {company_id})")
        
        # 2. PDF 파일을 Storage에 저장 (선택사항 - 에러 발생 시 무시)
        try:
            file_path = f"{company_id}/main.pdf"
            pdf_file.seek(0)
            pdf_bytes = pdf_file.read()
            supabase_client.storage.from_("company-pdfs").upload(
                file_path,
                pdf_bytes,
                {"content-type": "application/pdf"}
            )
            file_size = len(pdf_bytes)
            st.info("✅ PDF 파일 Storage 저장 완료")
        except Exception as storage_error:
            st.warning(f"⚠️ PDF Storage 저장 실패 (계속 진행): {storage_error}")
            file_path = "not_stored"
            file_size = 0
        
        # 3. PDF 파일 정보 저장
        pdf_data = {
            "company_id": company_id,
            "file_name": getattr(pdf_file, 'name', 'unknown.pdf'),
            "file_type": "main",
            "storage_path": file_path,
            "file_size": file_size,
            "extracted_text": extracted_text[:50000],  # 텍스트 크기 제한
            "pages_count": extracted_text.count("=== 페이지")
        }
        supabase_client.table("pdf_files").insert(pdf_data).execute()
        st.info("✅ PDF 메타데이터 저장 완료")
        
        # 4. 추출된 데이터 저장
        data_entries = []
        for field_name, field_value in extracted_data.items():
            data_entries.append({
                "company_id": company_id,
                "field_name": field_name,
                "field_value": str(field_value)[:5000]  # 길이 제한
            })
        
        if data_entries:
            supabase_client.table("extracted_data").insert(data_entries).execute()
            st.info(f"✅ 추출 데이터 {len(data_entries)}개 저장 완료")
        
        # 5. 보고서 저장 (선택사항)
        if report_content:
            report_data = {
                "company_id": company_id,
                "report_content": report_content[:100000]  # 크기 제한
            }
            supabase_client.table("reports").insert(report_data).execute()
            st.info("✅ 보고서 저장 완료")
        
        # 6. 구조화된 데이터 저장 (Upstage Parse 결과)
        if st.session_state.get('structured_data'):
            try:
                import json
                structured_json = json.dumps(st.session_state.structured_data, ensure_ascii=False)
                
                # companies 테이블에 structured_data 컬럼 추가 필요
                # 일단 extracted_data 테이블에 특수 필드로 저장
                supabase_client.table("extracted_data").insert({
                    "company_id": company_id,
                    "field_name": "__structured_data__",
                    "field_value": structured_json[:50000]  # 크기 제한
                }).execute()
                st.info("✅ 구조화된 데이터 저장 완료 (재사용 가능)")
            except Exception as e:
                st.warning(f"⚠️ 구조화된 데이터 저장 실패: {e}")
        
        # 7. 임베딩 생성 및 저장 (RAG 시스템)
        if create_embeddings_flag and openai_client:
            with st.spinner("🔮 임베딩 벡터 생성 중..."):
                # 텍스트 청크 분할
                chunks = split_text_into_chunks(extracted_text, max_tokens=500, overlap_tokens=50)
                st.info(f"📦 {len(chunks)}개 청크 생성 완료")
                
                # 임베딩 생성
                embeddings = create_embeddings(chunks)
                
                if embeddings:
                    # Supabase에 저장
                    save_embeddings_to_supabase(company_id, embeddings, file_type="main")
                else:
                    st.warning("⚠️ 임베딩 생성 실패 - 텍스트 검색은 제한됩니다")
        
        return company_id
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        st.error(f"❌ Supabase 저장 실패")
        st.error(f"에러: {str(e)}")
        with st.expander("상세 에러 로그"):
            st.code(error_detail)
        return None

def load_companies_list():
    """저장된 기업 목록 불러오기"""
    if not supabase_client:
        return []
    
    try:
        response = supabase_client.table("companies").select("*").order("created_at", desc=True).limit(50).execute()
        return response.data
    except Exception as e:
        st.error(f"기업 목록 로드 실패: {e}")
        return []

def load_company_data(company_id):
    """특정 기업의 추출된 데이터 및 구조화된 데이터 불러오기"""
    if not supabase_client:
        return {}, None
    
    try:
        response = supabase_client.table("extracted_data").select("*").eq("company_id", company_id).execute()
        
        extracted_data = {}
        structured_data = None
        
        for item in response.data:
            field_name = item["field_name"]
            field_value = item["field_value"]
            
            # 구조화된 데이터 복원
            if field_name == "__structured_data__":
                try:
                    import json
                    structured_data = json.loads(field_value)
                except:
                    pass
            else:
                extracted_data[field_name] = field_value
        
        return extracted_data, structured_data
    except Exception as e:
        st.error(f"데이터 로드 실패: {e}")
        return {}, None

# ============================================
# 임베딩 및 RAG 시스템
# ============================================

def split_text_into_chunks(text, max_tokens=500, overlap_tokens=50):
    """텍스트를 토큰 기반으로 청크 분할"""
    try:
        encoding = tiktoken.encoding_for_model("text-embedding-3-small")
        tokens = encoding.encode(text)
        
        chunks = []
        start = 0
        
        while start < len(tokens):
            end = start + max_tokens
            chunk_tokens = tokens[start:end]
            chunk_text = encoding.decode(chunk_tokens)
            
            chunks.append({
                "text": chunk_text,
                "start_pos": start,
                "end_pos": end,
                "token_count": len(chunk_tokens)
            })
            
            start += (max_tokens - overlap_tokens)
        
        return chunks
    except Exception as e:
        st.error(f"청크 분할 실패: {e}")
        # 폴백: 단순 문자 기반 분할
        chunk_size = 2000
        overlap = 200
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            chunks.append({
                "text": chunk_text,
                "start_pos": start,
                "end_pos": end,
                "token_count": len(chunk_text) // 4  # 대략적 추정
            })
            start += (chunk_size - overlap)
        return chunks

def create_embeddings(text_chunks):
    """OpenAI API로 임베딩 벡터 생성"""
    if not openai_client:
        st.error("OpenAI 클라이언트가 초기화되지 않았습니다.")
        return []
    
    embeddings = []
    try:
        for i, chunk in enumerate(text_chunks):
            response = openai_client.embeddings.create(
                model="text-embedding-3-small",
                input=chunk["text"]
            )
            embedding_vector = response.data[0].embedding
            
            embeddings.append({
                "chunk_index": i,
                "text": chunk["text"],
                "embedding": embedding_vector,
                "token_count": chunk["token_count"]
            })
            
            # 진행 상황 표시
            if (i + 1) % 10 == 0:
                st.info(f"임베딩 생성 중... {i + 1}/{len(text_chunks)}")
        
        return embeddings
    except Exception as e:
        st.error(f"임베딩 생성 실패: {e}")
        return []

def save_embeddings_to_supabase(company_id, embeddings, file_type="main"):
    """Supabase에 임베딩 벡터 저장"""
    if not supabase_client:
        st.warning("Supabase 클라이언트가 연결되지 않았습니다.")
        return False
    
    try:
        # 벡터 데이터 준비
        vector_entries = []
        for emb in embeddings:
            # 텍스트를 UTF-8로 명시적으로 인코딩/디코딩하여 한글 깨짐 방지
            chunk_text = emb["text"][:5000]
            # 이미 문자열이므로 그대로 사용 (Python 3는 기본 UTF-8)
            
            vector_entries.append({
                "company_id": company_id,
                "file_type": file_type,
                "chunk_index": emb["chunk_index"],
                "chunk_text": chunk_text,  # UTF-8 텍스트
                "embedding": emb["embedding"],
                "token_count": emb["token_count"]
            })
        
        # 배치로 저장 (한 번에 너무 많으면 분할)
        batch_size = 100
        for i in range(0, len(vector_entries), batch_size):
            batch = vector_entries[i:i + batch_size]
            supabase_client.table("document_embeddings").insert(batch).execute()
            st.info(f"벡터 저장 중... {min(i + batch_size, len(vector_entries))}/{len(vector_entries)}")
        
        st.success(f"✅ {len(vector_entries)}개 임베딩 벡터 저장 완료!")
        return True
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        st.error(f"벡터 저장 실패: {str(e)}")
        with st.expander("상세 에러 로그"):
            st.code(error_detail)
        return False

def semantic_search(query, company_id=None, top_k=5, file_type=None):
    """의미론적 유사도 검색"""
    if not supabase_client or not openai_client:
        st.warning("Supabase 또는 OpenAI 클라이언트가 연결되지 않았습니다.")
        return []
    
    try:
        # 쿼리 임베딩 생성
        response = openai_client.embeddings.create(
            model="text-embedding-3-small",
            input=query
        )
        query_embedding = response.data[0].embedding
        
        # Supabase에서 유사도 검색 (RPC 함수 사용)
        rpc_params = {
            "query_embedding": query_embedding,
            "match_threshold": 0.5,
            "match_count": top_k
        }
        
        # 회사 ID 필터
        if company_id:
            rpc_params["filter_company_id"] = company_id
        
        # 파일 타입 필터
        if file_type:
            rpc_params["filter_file_type"] = file_type
        
        # RPC 호출
        result = supabase_client.rpc(
            "match_documents",
            rpc_params
        ).execute()
        
        return result.data
    except Exception as e:
        st.error(f"유사도 검색 실패: {e}")
        return []

def retrieve_relevant_context(query, company_id=None, max_tokens=3000):
    """RAG: 쿼리와 관련된 컨텍스트 추출"""
    search_results = semantic_search(query, company_id=company_id, top_k=10)
    
    if not search_results:
        return "관련 컨텍스트를 찾을 수 없습니다."
    
    # 토큰 제한 내에서 관련 텍스트 조합
    context_parts = []
    total_tokens = 0
    
    for result in search_results:
        chunk_text = result.get("chunk_text", "")
        similarity = result.get("similarity", 0)
        token_count = result.get("token_count", 0)
        
        if total_tokens + token_count > max_tokens:
            break
        
        context_parts.append(f"[유사도: {similarity:.3f}]\n{chunk_text}")
        total_tokens += token_count
    
    return "\n\n---\n\n".join(context_parts)

# ============================================
# Upstage Document Parse API 연동
# ============================================
import requests

# Upstage API URL (최신 Document Digitization API)
UPSTAGE_API_URL = "https://api.upstage.ai/v1/document-digitization"

def check_upstage_available():
    """Upstage API 키 설정 확인"""
    return bool(UPSTAGE_API_KEY and UPSTAGE_API_KEY != "your-upstage-api-key-here")

# OCR Reader (lazy loading) - 로컬 폴백용
_ocr_reader = None

def get_ocr_reader():
    """OCR Reader를 lazy loading으로 가져오기 (로컬 폴백)"""
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr
        _ocr_reader = easyocr.Reader(['ko', 'en'], gpu=False)
    return _ocr_reader

def extract_text_from_pdf(pdf_file, max_pages=50, use_ocr=False):
    """PDF에서 텍스트 추출"""
    try:
        pdf_file.seek(0)
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        num_pages = min(len(doc), max_pages)
        
        text = ""
        for page_num in range(num_pages):
            page = doc[page_num]
            page_text = page.get_text()
            text += f"\n\n=== 페이지 {page_num+1} ===\n\n{page_text}"
        
        # 텍스트가 충분하고 OCR 요청 안 했으면 그대로 반환
        if len(text.strip()) > 100 and not use_ocr:
            doc.close()
            return text, num_pages
        
        # OCR 사용
        if use_ocr or len(text.strip()) < 100:
            if len(text.strip()) < 100:
                st.warning("텍스트 추출량이 적어 고급 분석을 사용합니다...")
            else:
                st.info("🔍 표 구조 인식 모드로 재추출합니다...")
            
            # Upstage API 시도
            if check_upstage_available():
                st.info("☁️ Upstage Document Parse 사용 (표 구조화 + OCR)")
                doc.close()
                pdf_file.seek(0)
                return extract_text_with_upstage(pdf_file, max_pages)
            else:
                # 로컬 OCR 폴백 (Upstage 없을 때만)
                st.warning("⚠️ Upstage API 미설정, 기본 OCR 사용")
                doc.close()
                pdf_file.seek(0)
                return extract_text_with_easyocr(pdf_file, max_pages)
        
        doc.close()
        return text, num_pages
        
    except Exception as e:
        st.error(f"PDF 읽기 오류: {e}")
        return "", 0

def extract_text_with_upstage(pdf_file, max_pages=50):
    """Upstage Document Parse API로 PDF 전체 분석 (표 구조화!)"""
    if not UPSTAGE_API_KEY:
        st.error("Upstage API 키가 설정되지 않았습니다.")
        return "", 0
    
    try:
        pdf_file.seek(0)
        pdf_bytes = pdf_file.read()
        
        st.info("🚀 Upstage Document Parse로 PDF 분석 중... (표 구조 인식)")
        
        # Upstage API 호출
        headers = {
            "Authorization": f"Bearer {UPSTAGE_API_KEY}"
        }
        
        files = {
            "document": (getattr(pdf_file, 'name', 'document.pdf'), pdf_bytes, "application/pdf")
        }
        
        # Upstage Document Parse API 파라미터 (표 + 차트 인식)
        data = {
            "ocr": "force",  # Always apply OCR
            "model": "document-parse",  # 명시적으로 모델 지정
            "output_formats": "['text', 'html', 'markdown']",  # JSON 배열을 문자열로
            "coordinates": "true",  # 좌표 정보 포함
            "base64_encoding": "['table', 'figure']",  # 표와 차트/그래프 모두 인코딩
        }
        
        response = requests.post(
            UPSTAGE_API_URL,
            headers=headers,
            files=files,
            data=data,
            timeout=120  # PDF 분석은 시간이 걸릴 수 있음
        )
        
        if response.status_code != 200:
            st.error(f"❌ Upstage API 오류: {response.status_code}")
            st.error(f"상세: {response.text}")
            return "", 0
        
        result = response.json()
        
        # 디버그: 응답 구조 확인
        with st.expander("🔍 디버그: Upstage API 응답 구조"):
            st.write("**응답 키:**", list(result.keys()))
            if "content" in result:
                st.write("**Content 키:**", list(result["content"].keys()))
            if "pages" in result:
                st.write(f"**페이지 수:** {len(result['pages'])}")
                if result['pages']:
                    first_page = result['pages'][0]
                    st.write("**첫 페이지 키:**", list(first_page.keys()))
                    if "elements" in first_page:
                        st.write(f"**첫 페이지 요소 수:** {len(first_page['elements'])}")
                        if first_page['elements']:
                            st.write("**첫 요소 예시:**", first_page['elements'][0])
        
        # 구조화된 텍스트 추출
        content = result.get("content", {})
        text = content.get("text", "")
        html = content.get("html", "")
        markdown = content.get("markdown", "")  # 마크다운도 추출
        
        # API v2.0 응답 구조: elements 배열에서 직접 추출
        elements_list = result.get("elements", [])
        
        # 구조화된 데이터 추출 (표, 차트, 제목, 리스트 등)
        structured_elements = {
            "tables": [],
            "charts": [],  # 차트/그래프 추가
            "headings": [],
            "paragraphs": [],
            "lists": []
        }
        
        # elements 배열에서 직접 추출 (v2.0 API)
        for element in elements_list:
            elem_category = element.get("category", "")
            elem_content = element.get("content", {})
            elem_page = element.get("page", 0)
            
            # content는 dict 형태 {html, markdown, text}
            elem_html = elem_content.get("html", "") if isinstance(elem_content, dict) else ""
            elem_text = elem_content.get("text", "") if isinstance(elem_content, dict) else str(elem_content)
            elem_markdown = elem_content.get("markdown", "") if isinstance(elem_content, dict) else ""
            
            if "table" in elem_category.lower():
                structured_elements["tables"].append({
                    "page": elem_page,
                    "content": elem_text or elem_html or elem_markdown,
                    "html": elem_html,
                    "markdown": elem_markdown
                })
            elif "figure" in elem_category.lower() or "chart" in elem_category.lower() or "image" in elem_category.lower():
                structured_elements["charts"].append({
                    "page": elem_page,
                    "content": elem_text or elem_html or elem_markdown,
                    "html": elem_html,
                    "markdown": elem_markdown,
                    "category": elem_category
                })
            elif "heading" in elem_category.lower() or "title" in elem_category.lower():
                structured_elements["headings"].append({
                    "page": elem_page,
                    "content": elem_text or elem_html
                })
            elif "list" in elem_category.lower():
                structured_elements["lists"].append({
                    "page": elem_page,
                    "content": elem_text or elem_html
                })
            elif "paragraph" in elem_category.lower():
                structured_elements["paragraphs"].append({
                    "page": elem_page,
                    "content": elem_text or elem_html
                })
        
        # 페이지별 정보 (호환성 유지)
        pages = result.get("pages", [])
        if not pages and elements_list:
            # pages가 없으면 elements로부터 생성
            num_pages = max([e.get("page", 1) for e in elements_list] + [1])
        else:
            num_pages = len(pages)
        
        # 표 및 차트 정보 추출
        table_count = len(structured_elements["tables"])
        chart_count = len(structured_elements["charts"])
        
        st.success(f"✅ Upstage 분석 완료: {num_pages}페이지, {len(text)}자, **표 {table_count}개, 차트 {chart_count}개** 인식")
        
        # 표가 인식되었을 때 상세 정보 표시
        if table_count > 0:
            with st.expander(f"📊 인식된 표 정보 ({table_count}개)"):
                for idx, table in enumerate(structured_elements["tables"][:3], 1):  # 최대 3개만 표시
                    st.write(f"**표 {idx} (페이지 {table.get('page', '?')})**")
                    table_content = table.get('html', '') or table.get('content', '')
                    if table_content:
                        st.text(table_content[:300] + ("..." if len(table_content) > 300 else ""))
                    st.markdown("---")
        
        # 차트가 인식되었을 때 상세 정보 표시
        if chart_count > 0:
            with st.expander(f"📈 인식된 차트/그래프 정보 ({chart_count}개)"):
                for idx, chart in enumerate(structured_elements["charts"][:3], 1):  # 최대 3개만 표시
                    st.write(f"**차트 {idx} (페이지 {chart.get('page', '?')}) - {chart.get('category', 'unknown')}**")
                    chart_content = chart.get('content', '') or chart.get('html', '')
                    if chart_content:
                        st.text(chart_content[:300] + ("..." if len(chart_content) > 300 else ""))
                    st.markdown("---")
        
        # 디버그: 표/차트 인식 실패 시 경고
        if table_count == 0 and chart_count == 0 and num_pages > 0:
            st.warning("⚠️ Upstage가 표와 차트를 인식하지 못했습니다.")
            st.info("**가능한 원인:**\n- PDF가 이미지 스캔본 (OCR 품질 저하)\n- 표/차트 구조가 복잡하거나 비정형\n- 텍스트로 된 표 형식 데이터")
            st.info("💡 **해결 방법:** LLM이 텍스트에서 직접 표 데이터를 추출하도록 프롬프트가 최적화되어 있습니다.")
            
            # 디버그 정보
            with st.expander("🔍 디버그: Upstage 응답 분석"):
                st.write("**API 응답 구조:**")
                st.write(f"- 전체 요소 수: {len(elements_list)}개")
                st.write(f"- 제목: {len(structured_elements['headings'])}개")
                st.write(f"- 단락: {len(structured_elements['paragraphs'])}개")
                st.write(f"- 리스트: {len(structured_elements['lists'])}개")
                st.write(f"- 표: {len(structured_elements['tables'])}개")
                st.write(f"- 차트: {len(structured_elements['charts'])}개")
                
                # elements 카테고리 분포
                categories = {}
                for elem in elements_list[:50]:  # 최대 50개
                    cat = elem.get("category", "unknown")
                    categories[cat] = categories.get(cat, 0) + 1
                
                if categories:
                    st.write("\n**요소 카테고리 분포:**")
                    for cat, count in sorted(categories.items(), key=lambda x: x[1], reverse=True):
                        st.write(f"  - {cat}: {count}개")
        
        # 세션에 구조화된 데이터 저장
        st.session_state.structured_data = structured_elements
        
        # 페이지별 텍스트 구조화 (elements로부터 재구성)
        structured_text = ""
        if elements_list:
            # elements를 페이지별로 그룹화
            pages_dict = {}
            for elem in elements_list:
                page_num = elem.get("page", 1)
                if page_num not in pages_dict:
                    pages_dict[page_num] = []
                
                content_obj = elem.get("content", {})
                if isinstance(content_obj, dict):
                    elem_text = content_obj.get("text", "") or content_obj.get("html", "")
                else:
                    elem_text = str(content_obj)
                
                if elem_text:
                    pages_dict[page_num].append(elem_text)
            
            # 페이지별로 텍스트 구성
            for page_num in sorted(pages_dict.keys())[:max_pages]:
                structured_text += f"\n\n=== 페이지 {page_num} ===\n\n"
                structured_text += "\n\n".join(pages_dict[page_num])
        
        # structured_text가 없으면 content.text 사용
        final_text = structured_text if structured_text.strip() else text
        
        return final_text, min(num_pages, max_pages)
        
    except requests.Timeout:
        st.error("⏱️ Upstage API 타임아웃 (대용량 PDF는 시간이 걸릴 수 있습니다)")
        return "", 0
    except Exception as e:
        st.error(f"❌ Upstage API 오류: {e}")
        import traceback
        st.error(traceback.format_exc())
        return "", 0

def extract_text_with_easyocr(pdf_file, max_pages=50):
    """로컬 EasyOCR로 텍스트 추출 (느림)"""
    text = ""
    try:
        pdf_file.seek(0)
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        num_pages = min(len(doc), max_pages)
        
        progress_bar = st.progress(0)
        
        for page_num in range(num_pages):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            
            if pix.n == 4:
                img_array = img_array[:, :, :3]
            
            ocr_reader = get_ocr_reader()
            ocr_result = ocr_reader.readtext(img_array, detail=0, paragraph=True)
            page_text = "\n".join(ocr_result)
            text += f"\n\n=== 페이지 {page_num+1} ===\n\n{page_text}"
            
            progress_bar.progress((page_num + 1) / num_pages)
        
        progress_bar.empty()
        doc.close()
        return text, num_pages
    except Exception as e:
        st.error(f"로컬 OCR 오류: {e}")
        return "", 0

def extract_all_keywords_batch(text, field_names, structured_data=None):
    """배치 방식으로 모든 키워드를 한 번에 추출 (구조화된 데이터 우선 활용)"""
    if not openai_client:
        # API 없으면 개별 방식으로 폴백
        result = {}
        for field_name in field_names:
            result[field_name] = extract_keyword_simple(text, field_name)
        return result
    
    try:
        # 구조화된 데이터가 있으면 우선 활용 - 표를 명확한 형식으로 변환
        context_info = ""
        has_structured_tables = False
        
        if structured_data:
            # 표 데이터를 마크다운/HTML 형식으로 변환
            if structured_data.get("tables"):
                has_structured_tables = True
                context_info += "\n\n" + "="*60 + "\n"
                context_info += "📊 **구조화된 표 데이터 (최우선 참조!)**\n"
                context_info += "="*60 + "\n\n"
                context_info += "⚠️ **재무 데이터는 아래 표에서만 추출하세요! 본문 텍스트 무시!**\n\n"
                
                for idx, table in enumerate(structured_data['tables']):  # 모든 표 표시
                    context_info += f"▶ **[표 {idx+1}] (페이지 {table.get('page', '?')})**\n\n"
                    
                    # Markdown이 가장 파싱하기 쉬우므로 우선
                    table_markdown = table.get('markdown', '')
                    table_html = table.get('html', '')
                    table_content = table.get('content', '')
                    
                    # 표 형식 선택 (Markdown > Content > HTML)
                    if table_markdown and len(table_markdown) > 20:
                        # Markdown 표를 더 명확하게 표시
                        context_info += "```표 (Markdown 형식)\n"
                        context_info += table_markdown[:1500]  # 증가
                        context_info += "\n```\n\n"
                    elif table_content and len(table_content) > 20:
                        # Content를 구조화해서 표시
                        context_info += "```표 (텍스트 형식)\n"
                        context_info += table_content[:1500]  # 증가
                        context_info += "\n```\n\n"
                    elif table_html:
                        context_info += "```표 (HTML 형식)\n"
                        context_info += table_html[:1000]
                        context_info += "\n```\n\n"
                    
                    # 표 해석 힌트 추가
                    context_info += "💡 이 표에서 행 이름(첫 번째 열)과 값들을 정확히 매칭하세요.\n\n"
                    
                context_info += "\n" + "="*60 + "\n"
                context_info += "🎯 **재무 데이터 추출 시 필수 확인:**\n"
                context_info += "- 표에서 '영업이익' 행을 찾고 해당 열의 숫자를 추출\n"
                context_info += "- 표에서 '매출액' 행을 찾고 해당 열의 숫자를 추출\n"
                context_info += "- 표 상단에 '단위: 억원' 같은 표시가 있으면 모든 숫자에 단위 적용\n"
                context_info += "- 본문에 '영업 생산성' 같은 비슷한 단어가 있어도 무시!\n"
                context_info += "="*60 + "\n\n"
            
            # 차트/그래프 데이터 추가
            if structured_data.get("charts"):
                context_info += "\n\n" + "="*60 + "\n"
                context_info += "📈 **인식된 차트/그래프 데이터**\n"
                context_info += "="*60 + "\n\n"
                
                for idx, chart in enumerate(structured_data['charts']):  # 모든 차트 표시
                    context_info += f"▶ **[차트 {idx+1}] (페이지 {chart.get('page', '?')}) - {chart.get('category', 'chart')}**\n\n"
                    
                    chart_content = chart.get('content', '') or chart.get('html', '')
                    if chart_content:
                        context_info += f"```\n{chart_content[:500]}\n```\n\n"
                
                context_info += "\n💡 차트 데이터에서 추출할 수 있는 정보(성장률, 추세 등)를 활용하세요.\n\n"
                context_info += "="*60 + "\n\n"
            
            # 주요 제목 요약 (문서 구조 파악용)
            if structured_data.get("headings"):
                context_info += "\n[📑 문서 구조 - 주요 섹션]\n"
                for heading in structured_data['headings'][:15]:  # 최대 15개
                    heading_text = heading['content'][:100]  # 긴 제목은 자르기
                    context_info += f"  • 페이지 {heading.get('page', '?')}: {heading_text}\n"
                context_info += "\n"
        
        # 텍스트 길이 조정 - 표가 있으면 중간, 없으면 길게
        if has_structured_tables:
            text_preview = text[:8000]  # 표가 있어도 충분한 컨텍스트 제공 (재무표 전체 포함)
        else:
            text_preview = text[:15000]  # 표가 없으면 텍스트에서 직접 찾아야 하므로 더 길게
        
        # 모든 필드를 한 번에 요청
        fields_list = "\n".join([f"{i+1}. {name}" for i, name in enumerate(field_names)])
        
        # 표가 있을 때와 없을 때 다른 프롬프트
        if has_structured_tables:
            extraction_guide = """
🎯 **추출 가이드 (구조화된 표 있음)**

1. **⭐ 표 데이터 절대 우선!**
   - 위에 제공된 "구조화된 표 데이터"를 **반드시 먼저** 분석하세요
   - 본문 텍스트는 보조 참고용으로만 사용하세요
   - HTML/Markdown 표 구조를 정확히 파싱하세요
   - 표의 헤더(열 이름)와 데이터 행을 구분하세요

2. **🔢 재무 데이터 추출 규칙 (엄격)**
   
   **💰 금액 데이터 (반드시 숫자 + 단위):**
   - 매출액: "Revenue", "Sales", "매출액" 행 → 예: "294억 원"
   - 영업이익: "Operating Profit", "영업이익" 행 → 예: "43억 원"
   - 순이익: "Net Profit", "Net Income", "당기순이익" 행 → 예: "24억 원"
   - EBITDA: "EBITDA" 행 찾기
   - CAPEX: "CAPEX", "자본적지출", "설비투자" 행
   - 현금흐름: "Cash Flow", "영업현금흐름", "OCF" 행
   - ⚠️ 주의: "17.6%"는 금액이 아닙니다! 비율/백분율은 제외!
   
   **📊 비율 데이터 (반드시 % 포함):**
   - 영업이익률: "Operating Margin", "영업이익률" 행 또는 (영업이익/매출액×100)
   - 순이익률: "Net Margin", "순이익률" 행 또는 (순이익/매출액×100)
   - 부채비율: "Debt Ratio", "부채비율" 행
   - ROE: "ROE", "자기자본이익률" 행
   - 성장률: "YoY", "Growth Rate", "CAGR" → 예: "+15.3%", "-9.3%"
   
   **⛔ 절대 금지:**
   - "영업 생산성", "영업 효율" 등은 영업이익이 아닙니다!
   - 본문에 "영업이익"이라는 단어가 있어도 표를 먼저 확인하세요!
   - 추측하거나 계산하지 마세요 (표에 직접 있는 값만!)

3. **🏢 기업 기본 정보 추출**
   - 회사명/기업명/법인명: 표지, 헤더, "회사명:" 라벨 찾기
   - 대표이사/CEO: "대표이사", "CEO", "Representative" 찾기
   - 설립일: "설립일", "설립연도", "Founded" (YYYY-MM-DD 또는 YYYY년)
   - 본사 위치: "본사", "Head Office", "Location", "주소"
   - 직원 수: "임직원수", "직원수", "Employees" (숫자+명)
   - 업종/산업분류: "업종", "Industry", "Sector"

4. **🏭 사업구조 & 제품 정보**
   - 사업분야: "사업영역", "Business Area", "주요사업" (여러 개면 쉼표로 구분)
   - 주요 제품: "제품", "Products", "Services" (구체적 제품명)
   - 핵심 기술: "Core Technology", "기술력", "R&D" (기술명)
   - 시장 점유율: "Market Share", "점유율" (% 또는 순위)
   - 고객사: "주요 고객", "거래처", "Customers" (기업명들)
   - 경쟁우위: "강점", "Competitive Advantage", "차별화"

5. **⚔️ 경쟁환경 & 리스크**
   - 경쟁사: "경쟁업체", "Competitors" (회사명들)
   - 시장 규모: "Market Size" (금액 + 단위)
   - 시장 성장률: "Market Growth Rate", "CAGR" (%)
   - 진입장벽: "Entry Barrier", "진입장벽"
   - SWOT 분석: "Strength", "Weakness", "Opportunity", "Threat"
   - 리스크: "Risk", "위험요인", "불확실성"
   - 규제 이슈: "규제", "Regulation"

6. **🚀 전략 & 미래 계획**
   - 신규 사업: "New Business", "신사업"
   - M&A: "인수합병", "M&A"
   - 투자 계획: "투자계획", "CAPEX", "설비투자"
   - 글로벌 진출: "해외진출", "Global Expansion"
   - R&D: "연구개발", "R&D 투자"
   - ESG 전략: "ESG", "지속가능경영", "탄소중립"

7. **📅 분기/연도 데이터 처리**
   - 분기별 데이터: "24.3Q", "25.2Q", "25.3Q" 등의 열
   - 최신 분기 데이터를 우선적으로 추출하세요
   - 여러 분기가 있으면 모두 나열: "43억 원 (25.3Q), 32억 원 (25.2Q)"

8. **📏 단위 인식 및 표기**
   - "단위: 억원" → 모든 숫자 뒤에 "억 원" 추가
   - "(십억 달러)" → "billion USD" 또는 "십억 달러"
   - "%", "비율" → 백분율 데이터
   - "명", "개", "건" → 개수 단위

9. **✅ 검증 체크리스트:**
   - [ ] 표에서 해당 키워드의 행을 찾았나요?
   - [ ] 금액은 숫자+단위 형태인가요? (예: 43억 원 ✅, 17.6% ❌)
   - [ ] 비율은 %가 포함되어 있나요?
   - [ ] 기업 정보는 정확하고 구체적인가요?
   - [ ] 여러 항목이 있으면 쉼표로 구분했나요?
   - [ ] 표에 없어서 본문을 봤다면, 정말 표에 없는 게 맞나요?
"""
        else:
            extraction_guide = """
🎯 **추출 가이드 (텍스트 기반 파싱)**

1. **📝 표 형식 텍스트 파싱**
   - "| 구분 | 24.3Q | 25.2Q |" → 표 헤더
   - "| 영업이익 | 561 | 390 |" → 데이터 행
   - 파이프(|) 구분자로 열을 나눠 파싱하세요
   - 표 형식이 여러 페이지에 걸쳐 있을 수 있으니 전체 스캔

2. **🔍 섹션별 탐색 우선순위**
   
   **재무 데이터:**
   ① "Financial Results", "경영실적", "영업실적", "손익계산서", "실적 요약"
   ② "재무정보", "재무현황", "재무상태표", "3Q Results"
   ③ 차트 제목 및 데이터 (Chart Type: bar/line 등)
   
   **기업 정보:**
   ① 첫 페이지, 표지, 헤더/푸터
   ② "Company Overview", "기업개요", "회사소개"
   ③ "Organization", "조직도"
   
   **사업 정보:**
   ① "Business", "사업구조", "사업영역"
   ② "Products & Services", "제품 및 서비스"
   ③ "Core Competency", "핵심역량"
   
   **전략 정보:**
   ① "Strategy", "전략", "Growth Strategy"
   ② "Future Plans", "향후 계획"
   ③ "Investment", "투자계획"
   
   ⚠️ 정보가 문서 전체에 분산되어 있을 수 있으니 전체를 스캔하세요

3. **🔢 재무 데이터 구별 (중요!)**
   
   **💰 금액 (숫자 + 단위 필수):**
   - 매출액: "2,345억 원", "294억 원" (✅)
   - 영업이익: "43억 원", "561억 원" (✅)
   - 순이익: "24억 원", "390억 원" (✅)
   - EBITDA: "450억 원" (✅)
   - CAPEX: "150억 원" (✅)
   - "17.6%"는 금액이 아닙니다! (❌)
   
   **📊 비율 (% 필수):**
   - 영업이익률: "14.6%", "23.5%" (✅)
   - 부채비율: "45.3%", "120%" (✅)
   - ROE: "15.2%" (✅)
   - 성장률: "+15.3%", "YoY -9.3%" (✅)
   
   **⛔ 혼동 주의:**
   - "영업 생산성 17.6%" ≠ 영업이익!
   - "영업 효율성 15%" ≠ 영업이익률!
   - "시장 점유율 25%" ≠ 성장률!

4. **🏢 기업 정보 추출 패턴**
   
   **회사명/기업명:**
   - "○○주식회사", "○○(주)", "○○ Co., Ltd."
   - 문서 상단, 로고 근처, "회사명:" 라벨
   
   **대표이사/CEO:**
   - "대표이사: 홍길동"
   - "CEO: John Doe"
   - "Representative Director"
   
   **설립일:**
   - "설립일: 1998년 3월 15일"
   - "Founded: 1998"
   - "Since 1998"
   
   **사업분야:**
   - "주요 사업: A, B, C"
   - "Business Areas: Manufacturing, Distribution"
   - 여러 개면 쉼표로 구분
   
   **주요 제품:**
   - 구체적인 제품명/서비스명
   - "제품 라인업:", "Product Portfolio:"
   
   **고객사:**
   - 기업명 나열: "삼성, LG, SK"
   - "Major Clients:", "주요 거래처:"

5. **⚔️ 경쟁 & 리스크 정보**
   
   **경쟁사:**
   - "경쟁업체:", "Competitors:"
   - 회사명들 나열
   
   **시장 규모:**
   - "시장 규모: 5조 원"
   - "Market Size: $5B"
   
   **SWOT 분석:**
   - "강점(Strength):", "약점(Weakness):"
   - "기회(Opportunity):", "위협(Threat):"
   
   **리스크:**
   - "리스크 요인:", "Risk Factors:"
   - "주요 위험:", "Risks:"

6. **🚀 전략 & 계획 정보**
   
   **신규 사업:**
   - "신사업:", "New Business:"
   - "사업 다각화", "Diversification"
   
   **M&A:**
   - "인수합병:", "M&A:"
   - "Acquisition", "Merger"
   
   **투자 계획:**
   - "투자 계획:", "Investment Plan:"
   - "CAPEX:", "설비투자:"
   
   **글로벌 진출:**
   - "해외 진출:", "Global Expansion:"
   - "수출:", "Export:"

7. **📋 패턴 매칭 예시**
   - "매출액: 2,345억 원" → "2,345억 원"
   - "영업이익률 23.5%" → "23.5%"
   - "24.3Q 영업이익 561억" → "561억 원 (24.3Q)"
   - "| 영업이익 | 43 | 32 |" → "43억 원 (최신), 32억 원"
   - "회사명: 동국생명과학" → "동국생명과학"
   - "대표이사: 홍길동" → "홍길동"
   - "주요 제품: A, B, C" → "A, B, C"

8. **🚫 실패 방지 전략**
   - 텍스트 전체를 꼼꼼히 스캔하세요
   - 유사 용어도 확인: 
     • "매출액" = "Sales" = "Revenue" = "총매출"
     • "영업이익" = "Operating Profit" = "Operating Income"
     • "순이익" = "Net Profit" = "Net Income" = "당기순이익"
     • "회사명" = "기업명" = "법인명" = "Company Name"
     • "대표이사" = "CEO" = "대표" = "Representative"
   - 약어도 확인:
     • "R&D" = "연구개발"
     • "M&A" = "인수합병"
     • "ESG" = "환경·사회·지배구조"
   - "정보 없음"은 정말 텍스트 어디에도 없을 때만!
   - 단위가 표시되어 있으면 반드시 포함하세요
   - 맥락에서 유추 가능한 정보도 활용하세요
"""
        
        prompt = f"""당신은 기업 실적 발표 자료 분석 전문가입니다. 아래 데이터에서 요청한 항목을 **정확히** 추출하세요.

{context_info}

{"[본문 텍스트 - 보조 참고용]" if has_structured_tables else "[본문 텍스트 - 주 분석 대상]"}
```
{text_preview}
```

---

**📋 추출할 항목:**
{fields_list}

{extraction_guide}

6. **⚠️ 출력 형식 (엄격히 준수!)**
   ```
   [항목명]: 추출된 값
   ```
   
   ✅ **올바른 예시 (모든 타입):**
   ```
   # 재무 데이터 (금액)
   [매출액]: 294억 원 (2025.3Q), 349억 원 (2025.2Q)
   [영업이익]: 43억 원 (2025.3Q), 32억 원 (2025.2Q)
   [순이익]: 24억 원 (2025.3Q), 29억 원 (2025.2Q)
   [EBITDA]: 450억 원 (2025.3Q)
   [CAPEX]: 150억 원 (2024년)
   [현금흐름]: 380억 원 (영업활동)
   
   # 재무 데이터 (비율)
   [영업이익률]: 14.6% (2025.3Q), 9.2% (2025.2Q)
   [순이익률]: 8.2% (2025.3Q)
   [부채비율]: 45.3% (2024년 말)
   [ROE]: 15.2% (2024년)
   [YoY]: -9.3% (매출액 기준)
   [CAGR]: +12.5% (2020-2024)
   
   # 기업 기본 정보
   [회사명]: 동국생명과학
   [기업명]: 동국생명과학 주식회사
   [대표이사]: 홍길동
   [CEO]: 홍길동
   [설립일]: 1998년 3월 15일
   [본사 위치]: 서울특별시 강남구
   [직원 수]: 350명 (2024년 기준)
   [업종]: 제약업
   [산업분류]: 의약품 제조업
   
   # 사업구조 & 제품
   [사업분야]: 조영제 제조 및 판매, 의료기기 유통, 헬스케어
   [주요 제품]: 파미레이, 메디레이, 유니레이, 가도비전
   [핵심 기술]: First Generic 기술력, 고순도 정제 기술, 수직 계열화
   [시장 점유율]: 국내 조영제 시장 21.4% (1위)
   [고객사]: 서울아산병원, 삼성서울병원, 세브란스병원 등 21개 상급병원
   [경쟁우위]: 국내 유일 수직 계열화, 최다 품목 라인업 43종
   
   # 경쟁환경
   [경쟁사]: A제약, B바이오, C헬스케어
   [시장 규모]: 국내 조영제 시장 5,000억 원 (2024년)
   [시장 성장률]: 연평균 7.5% 성장 (2020-2024)
   [진입장벽]: 높음 (인허가, 기술력, 유통망 필요)
   [SWOT 분석]: 강점-기술력/시장점유율, 약점-해외매출비중, 기회-고령화/진단수요, 위협-경쟁심화
   
   # 리스크
   [재무 리스크]: 환율 변동, 원재료 가격 상승
   [운영 리스크]: 품질 이슈, 생산 차질
   [규제 리스크]: 약가 인하 압력, 보험급여 정책 변화
   
   # 전략 & 미래
   [신규 사업]: AI 진단 소프트웨어 사업 진출 (2025년)
   [M&A]: 중소 의료기기 업체 인수 검토 중
   [투자 계획]: 2025년 CAPEX 200억 원 (생산설비 증설)
   [글로벌 진출]: 동남아시아 5개국 진출 (인도네시아, 베트남, 태국 등)
   [R&D]: 연간 매출의 5.5% R&D 투자 (신규 조영제 개발)
   [ESG 전략]: 2030년 탄소중립 달성 목표
   ```
   
   ❌ **잘못된 예시 (절대 금지!):**
   ```
   [영업이익]: 17.6% (← 이건 비율이지 금액이 아님!)
   [영업이익]: 영업 생산성 17.6% (← 영업이익 ≠ 영업 생산성!)
   [영업이익률]: 정보 없음 (← 표에 영업이익 43, 매출액 294가 있으면 계산 가능!)
   [매출액]: 약 300억 원 정도 (← 추측 금지! 정확한 값만!)
   [회사명]: 회사 (← 너무 불명확!)
   [대표이사]: CEO (← 이름을 찾아야 함!)
   [사업분야]: 제조업 (← 너무 일반적! 구체적으로!)
   [고객사]: 여러 병원 (← 구체적인 이름 필요!)
   ```

7. **🎯 필수 검증 체크리스트:**
   - [ ] 재무 데이터(매출, 이익 등)는 표에서 확인했나요?
   - [ ] 금액 항목에 숫자+단위(억 원, 달러)를 포함했나요?
   - [ ] 비율 항목(%로 끝나는 것)에 %를 포함했나요?
   - [ ] 본문의 "영업 생산성", "영업 효율" 등을 영업이익으로 착각하지 않았나요?
   - [ ] 기업명은 정확하고 공식 명칭인가요?
   - [ ] 사업분야는 구체적으로 작성했나요? (예: "제조업" ❌ → "의약품 제조업" ✅)
   - [ ] 여러 항목이 있으면 쉼표로 구분했나요?
   - [ ] 분기/연도 정보를 함께 표기했나요?
   - [ ] 정말 정보가 없어서 "정보 없음"이라고 했나요?

**🚀 지금 시작하세요! 표를 먼저 보고, 정확한 값을 추출하세요!**"""

        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "당신은 문서에서 정확한 정보를 추출하는 전문가입니다. 반드시 '[항목명]: 내용' 형식으로 답변합니다. 모든 요청된 항목에 대해 빠짐없이 답변합니다."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,  # 더 많은 키워드 처리 가능하도록 증가 (800 → 1500)
            temperature=0.1
        )
        
        result_text = response.choices[0].message.content.strip()
        print(f"[DEBUG] 배치 추출 결과:\n{result_text}\n")
        
        # 결과 파싱 - 더 간단한 방식
        extracted_data = {}
        
        for line in result_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # [항목명]: 또는 항목명: 형식 찾기
            if ':' in line:
                # [ ] 제거
                line = line.replace('[', '').replace(']', '')
                parts = line.split(':', 1)
                
                if len(parts) == 2:
                    field_name = parts[0].strip()
                    value = parts[1].strip()
                    
                    # 필드명이 요청한 항목 중 하나인지 확인
                    for fn in field_names:
                        if fn == field_name or fn in field_name or field_name in fn:
                            extracted_data[fn] = value
                            break
        
        # 누락된 필드는 "정보 없음"으로 채우기
        for field_name in field_names:
            if field_name not in extracted_data:
                extracted_data[field_name] = "정보 없음"
        
        print(f"[DEBUG] 파싱된 데이터: {extracted_data}\n")
        return extracted_data
        
    except Exception as e:
        print(f"[DEBUG] 배치 추출 실패: {e}")
        # 실패 시 개별 방식으로 폴백
        result = {}
        for field_name in field_names:
            result[field_name] = extract_keyword_simple(text, field_name)
        return result

def extract_keyword(text, field_name):
    """OpenAI API로 지능적으로 키워드 관련 정보 추출"""
    if not openai_client:
        return extract_keyword_simple(text, field_name)
    
    try:
        # 텍스트가 너무 길면 앞부분만 사용 (토큰 제한)
        text_preview = text[:4000]
        
        prompt = f"""다음 텍스트에서 "{field_name}"에 해당하는 정보를 찾아서 정확하게 추출하세요.

텍스트:
{text_preview}

요구사항:
1. "{field_name}"와 관련된 모든 정보를 추출
2. 정보가 없으면 "정보 없음"이라고만 응답
3. 추출한 정보를 그대로 답변 (설명 없이)
4. 여러 항목이 있으면 모두 포함
5. 원문의 표현을 최대한 유지

답변:"""

        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "당신은 문서에서 정확한 정보를 추출하는 전문가입니다. 요청받은 정보를 원문 그대로 완전하게 추출합니다."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=300,
            temperature=0.1
        )
        
        result = response.choices[0].message.content.strip()
        
        # 빈 응답이거나 너무 짧으면 폴백
        if not result or len(result) < 2:
            return extract_keyword_simple(text, field_name)
        
        return result
        
    except Exception as e:
        print(f"[DEBUG] OpenAI 추출 실패 ({field_name}): {e}")
        return extract_keyword_simple(text, field_name)

def extract_keyword_simple(text, field_name):
    """단순 텍스트 매칭 방식 (폴백)"""
    keywords_map = {
        "회사": ["회사", "기업", "법인", "㈜", "(주)", "주식회사"],
        "회사명": ["회사", "기업", "법인", "㈜", "(주)", "주식회사"],
        "회사이름": ["회사", "기업", "법인", "㈜", "(주)", "주식회사"],
        "기업명": ["회사", "기업", "법인", "㈜", "(주)", "주식회사"],
        "대표": ["대표", "CEO", "대표이사", "대표자"],
        "대표이름": ["대표", "CEO", "대표이사", "대표자"],
        "대표이사": ["대표", "CEO", "대표이사", "대표자"],
        "CEO": ["CEO", "대표", "대표이사"],
        "사업": ["사업", "업종", "사업분야", "사업내용"],
        "사업분야": ["사업분야", "사업", "업종", "주요 사업"],
        "사업내용": ["사업내용", "사업", "업종"],
        "매출": ["매출", "매출액", "Revenue"],
        "매출액": ["매출액", "매출", "Revenue"],
        "연매출": ["연매출", "매출액", "매출", "연간매출"],
        "영업이익": ["영업이익", "영업익", "Operating Profit"],
        "순이익": ["순이익", "당기순이익", "Net Profit"],
        "성과": ["성과", "실적", "주요 성과", "성과 지표"],
    }
    
    search_keywords = keywords_map.get(field_name, [field_name])
    
    lines = text.split('\n')
    result = []
    
    for keyword in search_keywords:
        pattern = re.compile(rf'{re.escape(keyword)}[:\s]*([^\n]+)', re.IGNORECASE)
        for line in lines:
            match = pattern.search(line)
            if match:
                value = match.group(1).strip()
                if value and len(value) > 1:
                    result.append(value)
    
    return result[0] if result else "정보 없음"

def generate_report_with_openai(data_dict, report_sections=None, model="gpt-4o-mini", company_id=None, use_rag=True, structured_data=None):
    """RAG 기반 + 구조화된 데이터 활용 OpenAI API로 체계적인 기업 분석 보고서 생성"""
    if not openai_client:
        return "❌ OpenAI API 키를 .env 파일에 설정하세요."
    
    # 보고서 섹션이 없으면 세션에서 가져오기
    if report_sections is None:
        report_sections = st.session_state.get('report_sections', [])
    
    # 선택된 섹션에 해당하는 작성 지침만 조합
    selected_guidelines = []
    for section in report_sections:
        if section in REPORT_SECTION_TEMPLATES:
            selected_guidelines.append(REPORT_SECTION_TEMPLATES[section])
    
    report_template = "\n\n".join(selected_guidelines)
    
    # 작성 규칙 추가
    report_template += """

**작성 규칙:**
- 각 섹션을 명확히 구분하여 작성
- PDF에서 추출한 정보는 그대로 사용
- PDF에 없는 정보로 외부 지식을 활용할 때는 반드시 "*> 본 문서에 해당 내용이 없어 외부 정보를 참고하여 답변 생성 (출처: [정보 출처])*" 형식으로 표시
- 마크다운 형식 사용 (## 제목, **강조**)
- 전문적이고 객관적인 톤 유지
- 구체적인 숫자와 데이터 중심으로 작성"""
    
    # 추출된 데이터 정리
    available_data = []
    missing_fields = []
    
    for key, value in data_dict.items():
        if value and value != "정보 없음":
            available_data.append(f"- {key}: {value}")
        else:
            missing_fields.append(key)
    
    if not available_data:
        return "❌ 추출된 데이터가 없습니다."
    
    available_data_text = "\n".join(available_data)
    missing_fields_text = ", ".join(missing_fields) if missing_fields else "없음"
    
    # RAG: 의미론적 검색으로 관련 컨텍스트 가져오기
    rag_context = ""
    if use_rag and company_id and supabase_client:
        with st.spinner("🔍 관련 문서 검색 중..."):
            # 보고서 섹션별 쿼리 생성
            queries = [
                f"{company_name} 재무 정보 매출 영업이익",
                f"{company_name} 사업 구조 제품 서비스",
                f"{company_name} 경쟁사 시장 분석",
                "리스크 요인 위험 요소"
            ]
            
            retrieved_contexts = []
            for query in queries:
                context = retrieve_relevant_context(query, company_id=company_id, max_tokens=1000)
                if context and context != "관련 컨텍스트를 찾을 수 없습니다.":
                    retrieved_contexts.append(context)
            
            if retrieved_contexts:
                rag_context = "\n\n**🔍 관련 문서 컨텍스트 (벡터 검색 결과):**\n" + "\n\n".join(retrieved_contexts[:2])  # 상위 2개만
                st.success(f"✅ {len(retrieved_contexts)}개 관련 컨텍스트 검색 완료")
    
    # 구조화된 데이터 컨텍스트 추가
    structured_context = ""
    if structured_data:
        structured_context = "\n\n**📊 문서 구조 정보 (Upstage Parse):**\n"
        
        # 표 데이터 요약
        if structured_data.get("tables"):
            structured_context += f"\n[표 데이터 {len(structured_data['tables'])}개 인식]\n"
            for idx, table in enumerate(structured_data['tables']):
                structured_context += f"\n표 {idx+1} (페이지 {table['page']}):\n{table['content'][:800]}\n"
        
        # 문서 구조
        if structured_data.get("headings"):
            structured_context += f"\n[문서 구조 - 주요 섹션]\n"
            for heading in structured_data['headings'][:15]:
                structured_context += f"- {heading['content']}\n"
    
    # 참고자료 텍스트 추가 (기존 방식)
    reference_context = ""
    if st.session_state.get('reference_pdfs'):
        reference_texts = []
        for filename, text in st.session_state.reference_pdfs.items():
            # 각 참고자료에서 앞부분 2000자만 사용 (토큰 절약)
            reference_texts.append(f"[{filename}]\n{text[:2000]}")
        
        reference_context = "\n\n**참고자료 (경쟁사/산업 분석 자료):**\n" + "\n\n---\n\n".join(reference_texts)
    
    try:
        prompt = f"""다음 기업 데이터를 바탕으로 전문적인 기업 분석 보고서를 작성하세요.

**PDF에서 추출된 데이터:**
{available_data_text}

**PDF에 없는 정보:** {missing_fields_text}

{structured_context}

{rag_context}

{reference_context}

**보고서 작성 지침:**
{report_template}

**중요: 표 데이터의 수치를 정확하게 인용하고, 문서 구조를 참고하여 체계적으로 작성하세요.**

**보고서:**"""
        
        response = openai_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": """당신은 기업 분석 전문가입니다. 
제공된 데이터를 바탕으로 체계적이고 전문적인 분석 보고서를 작성합니다.

**출처 표시 규칙 (필수):**
1. PDF에서 추출한 원본 데이터: `[출처: 메인 PDF]`
   예: "2023년 매출은 100억원입니다. [출처: 메인 PDF]"

2. 참고자료에서 가져온 정보: `[출처: 참고자료 - 파일명]`
   예: "경쟁사 A의 시장점유율은 30%입니다. [출처: 참고자료 - 산업분석.pdf]"

3. PDF 데이터를 분석/추론한 내용: `[분석: 메인 PDF 기반]`
   예: "영업이익률이 16%로 이자보상능력이 양호할 것으로 판단됩니다. [분석: 메인 PDF 기반]"

4. AI 학습 데이터 기반 일반 지식: `[출처: AI 학습 데이터 (2023년 10월 기준)]`
   예: "AI 산업은 기술 변화가 빠른 특성이 있습니다. [출처: AI 학습 데이터 (2023년 10월 기준)]"

5. PDF + AI 지식을 종합 분석: `[분석: 종합 판단]`
   예: "안정적 수익성을 보유하나 경쟁 심화로 모니터링이 필요합니다. [분석: 종합 판단]"

**중요:**
- 모든 문장에 반드시 출처를 명시하세요
- [추정], [예상] 같은 애매한 표현 사용 금지
- 출처가 명확하지 않으면 해당 내용을 작성하지 마세요

투자자와 대출 심사역이 읽기에 적합한 형식으로 작성합니다."""},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=2000
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        return f"❌ OpenAI API 오류: {str(e)}"

# 타이틀
st.markdown("""
<div style='
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    padding: 20px 40px;
    border-radius: 20px;
    margin-bottom: 30px;
    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.4);
    text-align: center;
'>
    <div style='
        font-size: 32px;
        font-weight: 800;
        color: white;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
        margin-bottom: 8px;
        letter-spacing: -1px;
    '>
        기업 분석 보고서 생성기
    </div>
    <div style='
        font-size: 15px;
        color: rgba(255, 255, 255, 0.95);
        font-weight: 500;
        letter-spacing: 0.5px;
    '>
        ✨ AI 기반 자동 정보 추출 및 전문 보고서 작성 시스템
    </div>
</div>
""", unsafe_allow_html=True)
st.markdown("")

# 사이드바 - 템플릿 설정
with st.sidebar:
    st.header("📝 템플릿 설정")
    
    # 이전 분석 불러오기
    if supabase_client:
        st.markdown("---")
        with st.expander("📂 이전 분석 불러오기"):
            companies = load_companies_list()
            if companies:
                company_names = [f"{c['company_name']} ({c['created_at'][:10]})" for c in companies]
                selected = st.selectbox("기업 선택", ["선택하세요..."] + company_names)
                
                if selected != "선택하세요...":
                    idx = company_names.index(selected)
                    company_id = companies[idx]["id"]
                    
                    if st.button("불러오기", type="primary"):
                        loaded_data, loaded_structured_data = load_company_data(company_id)
                        if loaded_data:
                            # 추출된 데이터 로드
                            st.session_state.extracted_data = loaded_data
                            
                            # 구조화된 데이터 로드 (Upstage Parse 결과)
                            if loaded_structured_data:
                                st.session_state.structured_data = loaded_structured_data
                                st.success(f"✅ 구조화된 데이터 복원 완료! (표 {len(loaded_structured_data.get('tables', []))}개)")
                            
                            # 템플릿 자동 생성 (키워드 복원)
                            st.session_state.template = []
                            for field_name in loaded_data.keys():
                                # 숫자 관련 키워드는 숫자 타입, 나머지는 텍스트 타입
                                field_type = "숫자" if any(keyword in field_name for keyword in ["매출", "이익", "비율", "YoY", "CAPEX", "ROE", "EBITDA", "부채", "현금"]) else "텍스트"
                                st.session_state.template.append({
                                    "name": field_name,
                                    "type": field_type
                                })
                            
                            st.success(f"✅ 데이터 로드 완료! ({len(loaded_data)}개 키워드)")
                            st.rerun()
            else:
                st.info("저장된 분석이 없습니다.")
        st.markdown("---")
    
    # 추천 키워드
    st.subheader("🎯 추천 키워드")
    
    categories = {
        "🏢 기업 기본 정보": [
            ("회사명", "텍스트"), ("기업명", "텍스트"), ("법인명", "텍스트"), ("설립일", "텍스트"),
            ("본사 위치", "텍스트"), ("대표이사", "텍스트"), ("CEO", "텍스트"), ("직원 수", "텍스트"),
            ("계열사", "텍스트"), ("업종", "텍스트"), ("산업분류", "텍스트"), ("기업 규모", "텍스트")
        ],
        "💰 재무정보": [
            ("매출액", "숫자"), ("영업이익", "숫자"), ("영업이익률", "숫자"), ("순이익", "숫자"),
            ("EBITDA", "숫자"), ("부채비율", "숫자"), ("현금흐름", "숫자"), ("ROE", "숫자"),
            ("CAPEX", "숫자"), ("전년 대비 증감", "숫자"), ("YoY", "숫자"), ("분기 실적", "숫자")
        ],
        "🏭 사업구조 & 제품": [
            ("사업분야", "텍스트"), ("주요 제품", "텍스트"), ("핵심 기술", "텍스트"), ("경쟁우위", "텍스트"),
            ("시장 점유율", "텍스트"), ("고객사", "텍스트"), ("유통 구조", "텍스트"), ("플랫폼", "텍스트")
        ],
        "⚔️ 경쟁환경": [
            ("경쟁사", "텍스트"), ("시장 규모", "텍스트"), ("시장 성장률", "텍스트"), ("진입장벽", "텍스트"),
            ("SWOT 분석", "텍스트"), ("규제 이슈", "텍스트"), ("산업 트렌드", "텍스트"), ("CAGR", "텍스트")
        ],
        "⚠️ 주요 리스크": [
            ("재무 리스크", "텍스트"), ("운영 리스크", "텍스트"), ("공급망 리스크", "텍스트"), ("기술 리스크", "텍스트"),
            ("규제 리스크", "텍스트"), ("환율 영향", "텍스트"), ("법적 이슈", "텍스트"), ("ESG 리스크", "텍스트")
        ],
        "🚀 기회 요인 & 전략": [
            ("신규 사업", "텍스트"), ("M&A", "텍스트"), ("투자 계획", "텍스트"), ("글로벌 진출", "텍스트"),
            ("R&D", "텍스트"), ("신제품 출시", "텍스트"), ("ESG 전략", "텍스트"), ("수익성 개선", "텍스트")
        ]
    }
    
    for category, keywords in categories.items():
        with st.expander(category):
            cols = st.columns(2)
            for idx, (kw, ftype) in enumerate(keywords):
                col = cols[idx % 2]
                if col.button(f"➕ {kw}", key=f"add_{category}_{kw}", use_container_width=True):
                    if not any(f['name'] == kw for f in st.session_state.template):
                        st.session_state.template.append({"name": kw, "type": ftype})
                        st.rerun()
    
    st.markdown("---")
    st.subheader("✏️ 직접 추가")
    
    with st.form("add_field_form"):
        field_name = st.text_input("필드 이름", placeholder="예: 회사이름")
        field_type = st.selectbox("데이터 타입", ["텍스트", "숫자"])
        submitted = st.form_submit_button("➕ 필드 추가", use_container_width=True)
        
        if submitted and field_name:
            if not any(f['name'] == field_name for f in st.session_state.template):
                st.session_state.template.append({"name": field_name, "type": field_type})
                st.rerun()
            else:
                st.warning(f"'{field_name}'은(는) 이미 추가되어 있습니다")
    
    col1, col2 = st.columns(2)
    if col1.button("📋 예시 로드", use_container_width=True):
        st.session_state.template = [
            {"name": "회사이름", "type": "텍스트"},
            {"name": "대표이름", "type": "텍스트"},
            {"name": "사업분야", "type": "텍스트"},
            {"name": "연매출", "type": "숫자"},
        ]
        st.rerun()
    
    if col2.button("🗑️ 초기화", use_container_width=True):
        st.session_state.template = []
        st.rerun()

# ============================================
# 메인 화면 - 사용자 로그인 체크
# ============================================

# 사용자 정보가 없으면 로그인 화면 표시
if not st.session_state.user_name:
    st.markdown("""
    <div style='text-align: center; padding: 50px 20px;'>
        <h1>🚀 기업 분석 보고서 생성기</h1>
        <p style='font-size: 18px; color: #666; margin-bottom: 40px;'>
            AI 기반 자동 분석 및 보고서 생성 시스템
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # 중앙 정렬된 로그인 폼
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("""
        <div style='background: white; padding: 40px; border-radius: 16px; 
        box-shadow: 0 10px 40px rgba(0,0,0,0.1); border: 2px solid #667eea;'>
        """, unsafe_allow_html=True)
        
        st.markdown("### 👤 사용자 정보 입력")
        st.info("📊 테스트 로그 수집을 위해 정보를 입력해주세요")
        
        with st.form("user_login_form"):
            user_name = st.text_input("이름 *", placeholder="홍길동", key="login_name")
            user_email = st.text_input("이메일 (선택)", placeholder="hong@example.com", key="login_email")
            
            col_btn1, col_btn2 = st.columns([1, 1])
            with col_btn1:
                submitted = st.form_submit_button("🚀 시작하기", use_container_width=True, type="primary")
            
            if submitted:
                if user_name:
                    st.session_state.user_name = user_name
                    st.session_state.user_email = user_email if user_email else None
                    
                    # 사용자 생성 또는 조회
                    if supabase_client:
                        user = create_or_get_test_user(user_name, user_email)
                        if user:
                            st.success(f"✅ {user_name}님 환영합니다!")
                            log_activity("user_login", "success", {"name": user_name, "email": user_email})
                            time.sleep(0.5)  # 성공 메시지 보여주기
                            st.rerun()
                    else:
                        st.session_state.user_name = user_name
                        st.session_state.user_email = user_email
                        st.success(f"✅ {user_name}님 환영합니다!")
                        time.sleep(0.5)
                        st.rerun()
                else:
                    st.error("❌ 이름을 입력해주세요")
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        # 하단 설명
        st.markdown("---")
        st.markdown("""
        ### ✨ 주요 기능
        - 📄 PDF 자동 분석 및 데이터 추출
        - 🤖 AI 기반 보고서 자동 생성
        - 💾 Supabase 데이터 저장 및 관리
        - 📊 실시간 로그 수집 및 모니터링
        """)
    
    st.stop()  # 로그인 전에는 아래 내용 표시 안 함

# ============================================
# 로그인 완료 후 - 사용자 정보 표시 (사이드바)
# ============================================
st.sidebar.markdown("---")
st.sidebar.markdown("### 👤 현재 사용자")
st.sidebar.success(f"✅ {st.session_state.user_name}님")
if st.session_state.user_email:
    st.sidebar.caption(f"📧 {st.session_state.user_email}")

if st.sidebar.button("🔄 다른 사용자로 변경", use_container_width=True):
    st.session_state.user_name = None
    st.session_state.user_email = None
    st.session_state.session_id = str(uuid.uuid4())
    st.session_state.current_test_session_id = None
    st.rerun()

# ============================================
# 메인 앱
# ============================================

# 메인 영역 - 관리자 탭은 특정 사용자에게만 표시
is_admin = (st.session_state.user_name == "신봉규" and 
            st.session_state.user_email == "shinbonggyu@daum.net")

if is_admin:
    tab1, tab2, tab3, tab_admin = st.tabs(["📋 템플릿 목록", "🔍 데이터 추출", "📄 보고서 생성", "🔧 관리자"])
else:
    tab1, tab2, tab3 = st.tabs(["📋 템플릿 목록", "🔍 데이터 추출", "📄 보고서 생성"])

with tab1:
    st.subheader("📋 현재 템플릿 목록")
    
    if st.session_state.template:
        st.markdown(f"**총 {len(st.session_state.template)}개 필드**")
        
        # 버튼 스타일로 표시 (추천 키워드처럼)
        cols = st.columns(4)  # 한 줄에 4개
        for idx, field in enumerate(st.session_state.template):
            col = cols[idx % 4]
            type_icon = "🔢" if field['type'] == "숫자" else "📝"
            
            with col:
                # 키워드명과 X 버튼을 한 줄로 표시
                button_col1, button_col2 = st.columns([4, 1])
                with button_col1:
                    st.markdown(f"""
                    <div style='
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        color: white;
                        padding: 10px 15px;
                        border-radius: 12px;
                        font-size: 13px;
                        font-weight: 500;
                        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                        margin-bottom: 10px;
                        text-align: left;
                    '>
                        {type_icon} {field['name']}
                    </div>
                    """, unsafe_allow_html=True)
                with button_col2:
                    if st.button("✕", key=f"delete_{idx}", help="삭제", use_container_width=True):
                        st.session_state.template.pop(idx)
                        st.rerun()
    else:
        st.info("템플릿이 비어있습니다. 왼쪽 사이드바에서 키워드를 추가하세요.")

with tab2:
    st.subheader("🔍 데이터 추출")
    st.info("PDF 파일을 업로드하고 AI가 자동으로 정보를 추출합니다")
    
    # Upstage API 상태 표시
    if check_upstage_available():
        st.success("🚀 Upstage Document Parse 연결됨 - 표 구조 인식 가능!")
    else:
        st.warning("⚠️ Upstage API 미설정 - 기본 텍스트 추출만 가능")
    
    # 메인 PDF 업로드
    st.markdown("### 📄 기업 보고서 (필수)")
    uploaded_file = st.file_uploader("기업 사업보고서 PDF 업로드", type=['pdf'], key="main_pdf")
    
    # 고급 분석 옵션 - 기본값을 True로 변경하고 강조
    st.markdown("---")
    st.markdown("### ⚙️ 추출 옵션")
    
    use_ocr_mode = st.checkbox(
        "📊 표 구조 인식 모드 (Upstage Document Parse) 🔥 권장",
        value=True,  # 기본값 True로 변경
        help="⭐ 재무제표, 실적 데이터가 포함된 PDF는 반드시 활성화하세요! 표를 인식하지 못하면 영업이익, 매출액 등을 추출할 수 없습니다."
    )
    
    if not use_ocr_mode:
        st.warning("⚠️ 표 구조 인식을 비활성화하면 재무 데이터 추출 성공률이 낮아집니다.")
    else:
        st.success("✅ 표 구조를 자동으로 인식하여 정확한 데이터 추출이 가능합니다!")
    
    # 참고자료 PDF 업로드 (RAG)
    st.markdown("---")
    st.markdown("### 📚 참고자료 추가 (선택)")
    st.info("💡 경쟁사 자료, 산업 리포트, 시장 분석 자료 등을 추가하면 더 정확한 보고서가 생성됩니다.")
    
    reference_files = st.file_uploader(
        "참고자료 PDF 업로드 (여러 개 가능)", 
        type=['pdf'], 
        accept_multiple_files=True,
        key="reference_pdfs_upload"
    )
    
    # 참고자료 처리
    if reference_files:
        st.markdown("**📋 업로드된 참고자료:**")
        for ref_file in reference_files:
            st.markdown(f"- {ref_file.name}")
    
    st.markdown("---")
    
    # 사용자 정보 확인
    if not st.session_state.user_name:
        st.warning("⚠️ 먼저 사이드바에서 사용자 정보를 입력하세요")
    elif uploaded_file and st.button("🚀 데이터 추출 시작", type="primary"):
        if not st.session_state.template:
            st.error("❌ 템플릿을 먼저 설정하세요! 사이드바에서 키워드를 추가하세요.")
        else:
            # 테스트 세션 시작
            session_start_time = time.time()
            company_name_temp = "Unknown"
            
            try:
                # 사용자 ID 가져오기
                user = create_or_get_test_user(st.session_state.user_name, st.session_state.user_email)
                user_id = user['id'] if user else None
                
                # 세션 시작 로그
                if user_id:
                    start_test_session(user_id, "처리 중", uploaded_file.name)
                
                # PDF 처리 시작
                log_activity("pdf_upload", "started", {"filename": uploaded_file.name, "ocr_mode": use_ocr_mode})
                
                with st.spinner("📄 PDF 처리 중..."):
                    pdf_start = time.time()
                    # 메인 PDF 텍스트 추출
                    pdf_text, num_pages = extract_text_from_pdf(uploaded_file, max_pages=50, use_ocr=use_ocr_mode)
                    st.session_state.pdf_text = pdf_text
                    pdf_time = int((time.time() - pdf_start) * 1000)
                    
                    log_activity("pdf_upload", "success", {
                        "filename": uploaded_file.name,
                        "pages": num_pages,
                        "text_length": len(pdf_text)
                    }, pdf_time)
                    
                    # 참고자료 PDF 처리
                    st.session_state.reference_pdfs = {}
                    if reference_files:
                        with st.spinner(f"📚 참고자료 {len(reference_files)}개 처리 중..."):
                            for ref_file in reference_files:
                                ref_text, ref_pages = extract_text_from_pdf(ref_file, max_pages=50, use_ocr=use_ocr_mode)
                                if ref_text:
                                    st.session_state.reference_pdfs[ref_file.name] = ref_text
                                    st.success(f"✅ {ref_file.name} 처리 완료 ({ref_pages}페이지, {len(ref_text)}자)")
                    
                    if pdf_text:
                        st.success(f"✅ 메인 PDF {num_pages}페이지 처리 완료 (총 {len(pdf_text)}자 추출)")
                        
                        # 배치 방식으로 키워드 추출 (구조화된 데이터 활용)
                        with st.spinner("🔍 데이터 추출 중..."):
                            extract_start = time.time()
                            field_names = [field['name'] for field in st.session_state.template]
                            
                            log_activity("keyword_extraction", "started", {"fields": field_names})
                            
                            # 구조화된 데이터가 있으면 전달
                            structured_data = st.session_state.get('structured_data')
                            if structured_data:
                                st.info(f"📊 구조화된 데이터 활용: 표 {len(structured_data.get('tables', []))}개, 제목 {len(structured_data.get('headings', []))}개")
                            
                            extracted_data = extract_all_keywords_batch(pdf_text, field_names, structured_data=structured_data)
                            st.session_state.extracted_data = extracted_data
                            extract_time = int((time.time() - extract_start) * 1000)
                            
                            # 회사명 추출
                            company_name_temp = extracted_data.get("회사명") or extracted_data.get("기업명") or "Unknown"
                            
                            log_activity("keyword_extraction", "success", {
                                "fields_count": len(field_names),
                                "extracted_count": len(extracted_data),
                                "company_name": company_name_temp
                            }, extract_time)
                            
                            # 🆕 데이터 품질 로그 기록 - OCR vs LLM 추출 비교
                            log_data_quality(
                                selected_keywords=field_names,
                                ocr_raw_text=pdf_text,
                                ocr_structured_data=structured_data,
                                llm_extracted_data=extracted_data,
                                llm_extraction_time_ms=extract_time,
                                company_name=company_name_temp,
                                pdf_filename=uploaded_file.name,
                                pdf_pages=num_pages
                            )
                        
                        # Supabase에 저장
                        if supabase_client:
                            with st.spinner("💾 Supabase에 저장 중..."):
                                save_start = time.time()
                                log_activity("data_save", "started")
                                
                                company_id = save_to_supabase(
                                    company_name=company_name_temp,
                                    pdf_file=uploaded_file,
                                    extracted_text=pdf_text,
                                    extracted_data=extracted_data
                                )
                                save_time = int((time.time() - save_start) * 1000)
                                
                                if company_id:
                                    st.success("✅ Supabase 저장 완료!")
                                    log_activity("data_save", "success", {"company_id": str(company_id)}, save_time)
                                else:
                                    log_activity("data_save", "failed", {"error": "company_id is None"}, save_time)
                        
                        # 세션 완료 로그
                        total_time = int((time.time() - session_start_time) * 1000)
                        complete_test_session("success", execution_time_ms=total_time)
                        
                        # 세션 업데이트 (회사명)
                        if st.session_state.current_test_session_id:
                            try:
                                supabase_client.table("test_sessions").update({
                                    "company_name": company_name_temp
                                }).eq("id", st.session_state.current_test_session_id).execute()
                            except:
                                pass
                        
                        # 결과 표시 - Gradio 스타일로
                        st.markdown("---")
                        st.markdown("## ✅ 처리 완료!")
                        st.markdown("### 🤖 AI가 자동으로 추출한 정보")
                        
                        # 추출된 데이터를 카드 형식으로 표시
                        for field in st.session_state.template:
                            value = extracted_data.get(field['name'], "정보 없음")
                            st.markdown(f"**📌 {field['name']}**")
                            st.markdown(f"""
                            <div style='padding: 10px; background: white; border-radius: 8px; 
                            margin-bottom: 15px; border: 1px solid #e2e8f0;'>
                            {value}
                            </div>
                            """, unsafe_allow_html=True)
                        
                        # 원본 텍스트 표시
                        st.markdown("---")
                        st.markdown("### 📄 추출된 원본 텍스트 (전체)")
                        st.markdown(f"""
                        <div style='background: #f8fafc; padding: 15px; border-radius: 8px; 
                        font-family: monospace; font-size: 13px; line-height: 1.6; 
                        max-height: 600px; overflow-y: auto;'>
                        {pdf_text}
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.error("❌ PDF에서 텍스트를 추출할 수 없습니다.")
                        log_error("pdf_upload", Exception("PDF 텍스트 추출 실패"))
                        complete_test_session("failed", "PDF 텍스트 추출 실패")
            
            except Exception as e:
                st.error(f"❌ 처리 중 오류 발생: {str(e)}")
                error_trace = traceback.format_exc()
                with st.expander("🔍 상세 에러 정보"):
                    st.code(error_trace)
                
                log_error("data_extraction", e, error_trace)
                complete_test_session("failed", str(e))
    
    # 이미 추출된 데이터가 있으면 표시 (새로 추출하거나 이전 분석 불러온 경우)
    elif st.session_state.extracted_data:
        st.markdown("---")
        
        # 불러온 데이터인지 확인
        if st.session_state.pdf_text:
            st.markdown("## ✅ 처리 완료!")
        else:
            st.markdown("## 📂 불러온 분석 데이터")
            st.info("💡 이전에 분석한 데이터를 불러왔습니다. 바로 보고서 생성이 가능합니다!")
            
            # 구조화된 데이터가 있으면 추가 키워드 추출 버튼 표시
            if st.session_state.get('structured_data'):
                st.markdown("### 🔄 추가 분석")
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.info("📊 구조화된 문서 데이터가 있습니다. 새로운 키워드를 템플릿에 추가하고 버튼을 눌러보세요!")
                with col2:
                    if st.button("➕ 추가 키워드 추출", type="primary"):
                        # 기존 템플릿과 추출된 데이터 비교
                        template_fields = {field['name'] for field in st.session_state.template}
                        existing_fields = set(st.session_state.extracted_data.keys())
                        new_fields = list(template_fields - existing_fields)
                        
                        if new_fields:
                            with st.spinner(f"🔍 {len(new_fields)}개 새 키워드 추출 중..."):
                                # PDF 텍스트가 없으면 구조화된 데이터에서 텍스트 재구성
                                if not st.session_state.pdf_text:
                                    structured_data = st.session_state.structured_data
                                    reconstructed_text = ""
                                    
                                    # 표 데이터 추가
                                    for table in structured_data.get('tables', []):
                                        reconstructed_text += f"\n{table['content']}\n"
                                    
                                    # 문단 데이터 추가
                                    for para in structured_data.get('paragraphs', [])[:50]:
                                        reconstructed_text += f"{para['content']}\n"
                                    
                                    st.session_state.pdf_text = reconstructed_text
                                
                                # 새 키워드만 추출
                                new_extracted = extract_all_keywords_batch(
                                    st.session_state.pdf_text,
                                    new_fields,
                                    structured_data=st.session_state.structured_data
                                )
                                
                                # 기존 데이터에 병합
                                st.session_state.extracted_data.update(new_extracted)
                                st.success(f"✅ {len(new_fields)}개 키워드 추출 완료! (API 비용 절감)")
                                st.rerun()
                        else:
                            st.warning("⚠️ 추출할 새 키워드가 없습니다. 템플릿에 키워드를 먼저 추가하세요!")
        
        st.markdown("### 🤖 추출된 정보")
        
        # 템플릿에 키워드가 있으면 템플릿 순서대로, 없으면 전체 표시
        if st.session_state.template:
            for field in st.session_state.template:
                value = st.session_state.extracted_data.get(field['name'], "정보 없음")
                st.markdown(f"**📌 {field['name']}**")
                st.markdown(f"""
                <div style='padding: 10px; background: white; border-radius: 8px; 
                margin-bottom: 15px; border: 1px solid #e2e8f0;'>
                {value}
                </div>
                """, unsafe_allow_html=True)
        else:
            # 템플릿 없으면 모든 데이터 표시
            for key, value in st.session_state.extracted_data.items():
                st.markdown(f"**📌 {key}**")
                st.markdown(f"""
                <div style='padding: 10px; background: white; border-radius: 8px; 
                margin-bottom: 15px; border: 1px solid #e2e8f0;'>
                {value}
                </div>
                """, unsafe_allow_html=True)
        
        # 구조화된 데이터 정보 표시 (표/차트)
        if st.session_state.get('structured_data'):
            structured_data = st.session_state.structured_data
            table_count = len(structured_data.get("tables", []))
            chart_count = len(structured_data.get("charts", []))
            
            st.markdown("---")
            st.markdown("### 📊 구조화된 데이터 분석 결과")
            
            # 표 정보
            if table_count > 0:
                with st.expander(f"📊 인식된 표 정보 ({table_count}개)", expanded=False):
                    for idx, table in enumerate(structured_data["tables"], 1):
                        st.write(f"**표 {idx} (페이지 {table.get('page', '?')})**")
                        table_content = table.get('html', '') or table.get('markdown', '') or table.get('content', '')
                        if table_content:
                            st.text(table_content[:300] + ("..." if len(table_content) > 300 else ""))
                        st.markdown("---")
            
            # 차트 정보
            if chart_count > 0:
                with st.expander(f"📈 인식된 차트/그래프 정보 ({chart_count}개)", expanded=False):
                    for idx, chart in enumerate(structured_data["charts"], 1):
                        st.write(f"**차트 {idx} (페이지 {chart.get('page', '?')}) - {chart.get('category', 'unknown')}**")
                        chart_content = chart.get('content', '') or chart.get('html', '')
                        if chart_content:
                            st.text(chart_content[:300] + ("..." if len(chart_content) > 300 else ""))
                        st.markdown("---")
        
        # 원본 텍스트가 있을 때만 표시 (새로 추출한 경우)
        if st.session_state.pdf_text:
            st.markdown("---")
            st.markdown("### 📄 추출된 원본 텍스트 (전체)")
            st.markdown(f"""
            <div style='background: #f8fafc; padding: 15px; border-radius: 8px; 
            font-family: monospace; font-size: 13px; line-height: 1.6; 
            max-height: 600px; overflow-y: auto;'>
            {st.session_state.pdf_text}
            </div>
            """, unsafe_allow_html=True)

with tab3:
    st.subheader("📄 보고서 생성")
    st.info("AI가 추출한 데이터를 바탕으로 전문 보고서를 자동 생성합니다")
    
    if not st.session_state.extracted_data:
        st.warning("⚠️ 먼저 '데이터 추출' 탭에서 PDF 데이터를 추출하세요.")
    else:
        st.markdown("### 📊 추출된 데이터 확인")
        
        # 템플릿 기준으로 표시 (템플릿에 있는 키워드만)
        if st.session_state.template:
            displayed_count = 0
            for field in st.session_state.template:
                field_name = field['name']
                if field_name in st.session_state.extracted_data:
                    value = st.session_state.extracted_data[field_name]
                    st.markdown(f"**{field_name}**: {value}")
                    displayed_count += 1
            
            # 템플릿에 없는 키워드도 표시할지 확인
            all_keys = set(st.session_state.extracted_data.keys())
            template_keys = {field['name'] for field in st.session_state.template}
            extra_keys = all_keys - template_keys
            
            if extra_keys:
                with st.expander(f"➕ 템플릿에 없는 추가 데이터 ({len(extra_keys)}개)"):
                    for key in extra_keys:
                        st.markdown(f"**{key}**: {st.session_state.extracted_data[key]}")
        else:
            # 템플릿이 없으면 모든 데이터 표시
            for key, value in st.session_state.extracted_data.items():
                st.markdown(f"**{key}**: {value}")
        
        st.markdown("---")
        
        col1, col2, col3 = st.columns([2, 2, 1])
        
        with col1:
            if st.button("📋 보고서 미리보기", type="secondary"):
                report_start = time.time()
                log_activity("report_generation", "started", {
                    "sections_count": len(st.session_state.get('report_sections', []))
                })
                
                with st.spinner("✨ OpenAI로 보고서 생성 중..."):
                    try:
                        # 구조화된 데이터 전달
                        structured_data = st.session_state.get('structured_data')
                        
                        report = generate_report_with_openai(
                            data_dict=st.session_state.extracted_data,
                            report_sections=st.session_state.get('report_sections'),
                            structured_data=structured_data
                        )
                        
                        # 참고자료 정보 추가
                        if st.session_state.get('reference_pdfs'):
                            ref_list = list(st.session_state.reference_pdfs.keys())
                            report += f"\n\n---\n\n**📚 참고자료 목록:**\n"
                            for ref_file in ref_list:
                                report += f"- {ref_file}\n"
                        
                        # 구조 정보 추가
                        if structured_data:
                            report += f"\n\n**📊 문서 분석 정보 (Upstage Parse):**\n"
                            report += f"- 표 {len(structured_data.get('tables', []))}개 인식\n"
                            report += f"- 섹션 {len(structured_data.get('headings', []))}개 구조화\n"
                        
                        # 보고서를 세션에 저장
                        st.session_state.report = report
                        
                        report_time = int((time.time() - report_start) * 1000)
                        log_activity("report_generation", "success", {
                            "report_length": len(report),
                            "sections": st.session_state.get('report_sections', [])
                        }, report_time)
                        
                        # 🆕 데이터 품질 로그 업데이트 - 보고서 생성 데이터 추가
                        if st.session_state.current_test_session_id:
                            # 기존 로그를 찾아서 업데이트
                            try:
                                logs = supabase_client.table("data_quality_logs")\
                                    .select("*")\
                                    .eq("session_id", st.session_state.current_test_session_id)\
                                    .order("created_at", desc=True)\
                                    .limit(1)\
                                    .execute()
                                
                                if logs.data and len(logs.data) > 0:
                                    latest_log = logs.data[0]
                                    supabase_client.table("data_quality_logs").update({
                                        "report_generated": True,
                                        "report_content": report[:20000],  # 처음 20000자만 저장
                                        "report_model": "gpt-4o-mini",
                                        "report_generation_time_ms": report_time
                                    }).eq("id", latest_log['id']).execute()
                                    print(f"✅ 데이터 품질 로그 업데이트 완료: {latest_log['id']}")
                            except Exception as e:
                                print(f"⚠️ 데이터 품질 로그 업데이트 실패: {e}")
                        
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ 보고서 생성 실패: {str(e)}")
                        error_trace = traceback.format_exc()
                        log_error("report_generation", e, error_trace)
                        with st.expander("🔍 상세 에러 정보"):
                            st.code(error_trace)
        
        with col2:
            if st.button("📄 보고서 생성 (DOCX)", type="primary"):
                if 'report' not in st.session_state or not st.session_state.report:
                    st.error("❌ 먼저 '보고서 미리보기' 버튼으로 보고서를 생성하세요.")
                else:
                    try:
                        # DOCX 생성
                        doc = Document()
                        
                        # 제목
                        title = doc.add_heading('기업 분석 보고서', 0)
                        title.alignment = 1  # 중앙 정렬
                        
                        # 작성일
                        date_para = doc.add_paragraph(f'작성일: {datetime.now().strftime("%Y년 %m월 %d일")}')
                        date_para.alignment = 1
                        doc.add_paragraph('')
                        
                        # 보고서 내용 파싱 및 추가
                        lines = st.session_state.report.split('\n')
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                            
                            # 제목 처리 (## 시작)
                            if line.startswith('## '):
                                doc.add_heading(line.replace('## ', ''), 1)
                            elif line.startswith('### '):
                                doc.add_heading(line.replace('### ', ''), 2)
                            # 볼드 처리 (**텍스트**)
                            elif line.startswith('**') and line.endswith('**'):
                                p = doc.add_paragraph()
                                p.add_run(line.strip('*')).bold = True
                            # 리스트 처리
                            elif line.startswith('- ') or line.startswith('* '):
                                doc.add_paragraph(line[2:], style='List Bullet')
                            # 일반 문단
                            else:
                                doc.add_paragraph(line)
                        
                        doc.add_paragraph('')
                        doc.add_paragraph('─' * 50)
                        doc.add_paragraph('')
                        
                        # 추출된 데이터 테이블
                        doc.add_heading('📋 추출된 상세 데이터', 1)
                        
                        table = doc.add_table(rows=1, cols=2)
                        table.style = 'Light Grid Accent 1'
                        
                        hdr = table.rows[0].cells
                        hdr[0].text = '항목'
                        hdr[1].text = '내용'
                        
                        for key, val in st.session_state.extracted_data.items():
                            row = table.add_row().cells
                            row[0].text = key
                            row[1].text = str(val)
                        
                        # 파일 저장
                        output_path = "기업_분석_보고서.docx"
                        doc.save(output_path)
                        
                        # 다운로드 버튼 제공
                        with open(output_path, "rb") as file:
                            st.download_button(
                                label="📥 보고서 다운로드",
                                data=file,
                                file_name=f"기업_분석_보고서_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        st.success("✅ DOCX 보고서가 생성되었습니다!")
                        
                    except Exception as e:
                        st.error(f"❌ DOCX 생성 실패: {str(e)}")
        
        with col3:
            if st.button("⚙️ 템플릿", help="보고서 작성 지침 수정"):
                st.session_state.show_template_editor = not st.session_state.show_template_editor
                st.rerun()
        
        # 보고서 템플릿 편집기
        if st.session_state.show_template_editor:
            st.markdown("---")
            st.markdown("### ⚙️ 보고서 섹션 선택")
            st.info("💡 원하는 보고서 섹션을 선택하세요. 선택한 섹션만 보고서에 포함됩니다.")
            
            # 모든 섹션 체크박스
            available_sections = [
                "기업 개요",
                "사업 구조 및 Revenue Model 분석",
                "산업 및 시장 분석",
                "재무 요약",
                "재무 건전성 심화 분석",
                "고객사 및 매출 집중도 분석",
                "경쟁사 비교 분석",
                "경영진 역량 및 지배구조 분석",
                "신용도 및 법률 리스크",
                "리스크 요인",
                "종합 평가"
            ]
            
            # 2열로 배치
            cols = st.columns(2)
            selected_sections = []
            
            for idx, section in enumerate(available_sections):
                col = cols[idx % 2]
                with col:
                    is_selected = st.checkbox(
                        section,
                        value=section in st.session_state.report_sections,
                        key=f"section_{section}"
                    )
                    if is_selected:
                        selected_sections.append(section)
            
            col_save, col_reset = st.columns(2)
            with col_save:
                if st.button("💾 선택 저장", type="primary", use_container_width=True):
                    st.session_state.report_sections = selected_sections
                    st.success(f"✅ {len(selected_sections)}개 섹션이 선택되었습니다!")
            
            with col_reset:
                if st.button("🔄 전체 선택", use_container_width=True):
                    st.session_state.report_sections = available_sections.copy()
                    st.success("✅ 모든 섹션이 선택되었습니다!")
                    st.rerun()
            
            # 선택된 섹션 미리보기
            if selected_sections:
                with st.expander("📋 선택된 섹션 미리보기", expanded=False):
                    for section in selected_sections:
                        st.markdown(f"**✓ {section}**")
                        if section in REPORT_SECTION_TEMPLATES:
                            st.text(REPORT_SECTION_TEMPLATES[section])
                        st.markdown("---")
        
        # 생성된 보고서가 있으면 항상 표시
        if 'report' in st.session_state and st.session_state.report:
            st.markdown("---")
            st.markdown("### 📄 생성된 보고서")
            st.markdown(st.session_state.report)

# ============================================
# 관리자 페이지
# ============================================
if is_admin:
    with tab_admin:
        st.subheader("🔧 관리자 페이지")
        
        # 비밀번호 확인
        if 'admin_logged_in' not in st.session_state:
            st.session_state.admin_logged_in = False
        
        if not st.session_state.admin_logged_in:
            st.info("🔒 관리자 페이지는 비밀번호가 필요합니다")
            admin_password = st.text_input("비밀번호", type="password", key="admin_password")
            
            if st.button("로그인"):
                # 환경변수에서 비밀번호 가져오기 (기본값: admin123)
                correct_password = os.getenv("ADMIN_PASSWORD", "admin123")
                if admin_password == correct_password:
                    st.session_state.admin_logged_in = True
                    st.rerun()
                else:
                    st.error("❌ 비밀번호가 틀렸습니다")
        else:
            st.success("✅ 관리자 로그인됨")
            
            if st.button("🚪 로그아웃"):
                st.session_state.admin_logged_in = False
                st.rerun()
            
            st.markdown("---")
            
            if not supabase_client:
                st.warning("⚠️ Supabase가 연결되지 않아 로그를 조회할 수 없습니다")
            else:
                # 탭 구성
                admin_tab1, admin_tab2, admin_tab3, admin_tab4 = st.tabs(["📊 통계", "👥 사용자 목록", "📋 로그 조회", "🔍 데이터 품질 비교"])
                
                with admin_tab1:
                    st.markdown("### 📊 테스트 통계")
                    
                    try:
                        # 전체 세션 수
                        sessions = supabase_client.table("test_sessions").select("*").execute()
                        total_sessions = len(sessions.data) if sessions.data else 0
                        
                        # 성공/실패 세션
                        success_sessions = len([s for s in sessions.data if s.get('status') == 'success']) if sessions.data else 0
                        failed_sessions = len([s for s in sessions.data if s.get('status') == 'failed']) if sessions.data else 0
                        in_progress = len([s for s in sessions.data if s.get('status') == 'in_progress']) if sessions.data else 0
                        
                        # 사용자 수
                        users = supabase_client.table("test_users").select("*").execute()
                        total_users = len(users.data) if users.data else 0
                        
                        # 메트릭 표시
                        col1, col2, col3, col4 = st.columns(4)
                        col1.metric("👥 총 사용자", total_users)
                        col2.metric("📝 총 세션", total_sessions)
                        col3.metric("✅ 성공", success_sessions)
                        col4.metric("❌ 실패", failed_sessions)
                        
                        if in_progress > 0:
                            st.info(f"⏳ 진행 중인 세션: {in_progress}개")
                        
                        # 성공률
                        if total_sessions > 0:
                            success_rate = (success_sessions / total_sessions) * 100
                            st.progress(success_rate / 100)
                            st.caption(f"성공률: {success_rate:.1f}%")
                        
                    except Exception as e:
                        st.error(f"통계 조회 실패: {e}")
                
                with admin_tab2:
                    st.markdown("### 👥 사용자 목록")
                    
                    try:
                        users = supabase_client.table("test_users").select("*").order("created_at", desc=True).execute()
                        
                        if users.data:
                            for user in users.data:
                                with st.expander(f"👤 {user.get('name', 'Unknown')} ({user.get('email', 'N/A')})"):
                                    st.write(f"**세션 ID**: `{user.get('session_id', 'N/A')}`")
                                    st.write(f"**가입일**: {user.get('created_at', 'N/A')}")
                                    
                                    # 해당 사용자의 세션 조회
                                    user_sessions = supabase_client.table("test_sessions").select("*").eq("user_id", user['id']).order("started_at", desc=True).execute()
                                    
                                    if user_sessions.data:
                                        st.write(f"**총 세션 수**: {len(user_sessions.data)}")
                                        for session in user_sessions.data[:5]:  # 최근 5개만
                                            status_emoji = "✅" if session.get('status') == 'success' else "❌" if session.get('status') == 'failed' else "⏳"
                                            st.write(f"{status_emoji} {session.get('company_name', 'N/A')} - {session.get('started_at', 'N/A')}")
                        else:
                            st.info("등록된 사용자가 없습니다")
                    
                    except Exception as e:
                        st.error(f"사용자 목록 조회 실패: {e}")
                
                with admin_tab3:
                    st.markdown("### 📋 로그 조회 및 다운로드")
                    
                    # 필터
                    col1, col2 = st.columns(2)
                    with col1:
                        log_type = st.selectbox("로그 유형", ["전체", "세션 로그", "활동 로그", "에러만"])
                    with col2:
                        limit = st.number_input("표시 개수", 10, 500, 100)
                    
                    if st.button("🔍 로그 조회", type="primary"):
                        try:
                            if log_type == "세션 로그" or log_type == "전체":
                                st.markdown("#### 📝 세션 로그")
                                sessions = supabase_client.table("test_sessions").select("*").order("started_at", desc=True).limit(limit).execute()
                                
                                if sessions.data:
                                    for session in sessions.data:
                                        status_color = "green" if session.get('status') == 'success' else "red" if session.get('status') == 'failed' else "orange"
                                        st.markdown(f"**:{status_color}[{session.get('status', 'unknown').upper()}]** {session.get('company_name', 'N/A')} - {session.get('pdf_filename', 'N/A')}")
                                        st.caption(f"시작: {session.get('started_at', 'N/A')} | 완료: {session.get('completed_at', 'N/A')}")
                                        if session.get('error_message'):
                                            with st.expander("❌ 에러 메시지"):
                                                st.code(session.get('error_message'))
                                        st.markdown("---")
                                    
                                    # CSV 다운로드
                                    import pandas as pd
                                    df = pd.DataFrame(sessions.data)
                                    csv = df.to_csv(index=False, encoding='utf-8-sig')
                                    st.download_button(
                                        "📥 세션 로그 CSV 다운로드",
                                        csv,
                                        f"session_logs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                        "text/csv"
                                    )
                            
                            if log_type == "활동 로그" or log_type == "전체":
                                st.markdown("#### 🔍 활동 로그")
                                
                                query = supabase_client.table("activity_logs").select("*").order("created_at", desc=True).limit(limit)
                                if log_type == "에러만":
                                    query = query.eq("status", "failed")
                                
                                logs = query.execute()
                                
                                if logs.data:
                                    for log in logs.data:
                                        status_emoji = "✅" if log.get('status') == 'success' else "❌" if log.get('status') == 'failed' else "⏳"
                                        st.markdown(f"{status_emoji} **{log.get('step', 'unknown')}** - {log.get('status', 'unknown')}")
                                        st.caption(f"시간: {log.get('created_at', 'N/A')} | 실행시간: {log.get('execution_time_ms', 0)}ms")
                                        
                                        if log.get('details'):
                                            with st.expander("📄 상세 정보"):
                                                st.json(log.get('details'))
                                        st.markdown("---")
                                    
                                    # CSV 다운로드
                                    import pandas as pd
                                    df = pd.DataFrame(logs.data)
                                    csv = df.to_csv(index=False, encoding='utf-8-sig')
                                    st.download_button(
                                        "📥 활동 로그 CSV 다운로드",
                                        csv,
                                        f"activity_logs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                        "text/csv"
                                    )
                                else:
                                    st.info("조회된 로그가 없습니다")
                        
                        except Exception as e:
                            st.error(f"로그 조회 실패: {e}")
                            st.code(traceback.format_exc())
                    
                    st.markdown("---")
                    st.markdown("### 📦 전체 로그 다운로드")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("📥 세션 로그 전체 다운로드"):
                            try:
                                sessions = supabase_client.table("test_sessions").select("*").order("started_at", desc=True).execute()
                                if sessions.data:
                                    import pandas as pd
                                    df = pd.DataFrame(sessions.data)
                                    csv = df.to_csv(index=False, encoding='utf-8-sig')
                                    st.download_button(
                                        "다운로드",
                                        csv,
                                        f"all_sessions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                        "text/csv",
                                        key="download_all_sessions"
                                    )
                            except Exception as e:
                                st.error(f"다운로드 실패: {e}")
                    
                    with col2:
                        if st.button("📥 활동 로그 전체 다운로드"):
                            try:
                                logs = supabase_client.table("activity_logs").select("*").order("created_at", desc=True).limit(5000).execute()
                                if logs.data:
                                    import pandas as pd
                                    df = pd.DataFrame(logs.data)
                                    csv = df.to_csv(index=False, encoding='utf-8-sig')
                                    st.download_button(
                                        "다운로드",
                                        csv,
                                        f"all_logs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                        "text/csv",
                                        key="download_all_logs"
                                    )
                            except Exception as e:
                                st.error(f"다운로드 실패: {e}")                
                with admin_tab4:
                    st.markdown("### 🔍 데이터 품질 비교 분석")
                    st.info("📊 OCR 추출 → LLM 데이터 추출 → 보고서 생성 과정을 비교하여 데이터 품질을 검증합니다")
                    
                    try:
                        # 데이터 품질 로그 조회
                        quality_logs = supabase_client.table("data_quality_logs")\
                            .select("*")\
                            .order("created_at", desc=True)\
                            .limit(50)\
                            .execute()
                        
                        if not quality_logs.data:
                            st.warning("아직 데이터 품질 로그가 없습니다. 데이터 추출을 먼저 진행하세요.")
                        else:
                            # 로그 목록 표시
                            st.markdown(f"**총 {len(quality_logs.data)}개의 품질 로그**")
                            
                            # 로그 선택
                            log_options = []
                            for log in quality_logs.data:
                                created_at = log.get('created_at', 'N/A')[:19].replace('T', ' ')
                                company = log.get('company_name', 'Unknown')
                                user = log.get('user_name', 'N/A')
                                keywords_count = len(log.get('selected_keywords', []))
                                success_rate = log.get('extraction_success_rate', 0)
                                report_gen = "✅ 보고서 있음" if log.get('report_generated') else "❌ 보고서 없음"
                                
                                log_options.append(
                                    f"{created_at} | {company} | {user} | 키워드 {keywords_count}개 | 성공률 {success_rate}% | {report_gen}"
                                )
                            
                            selected_log_idx = st.selectbox(
                                "분석할 로그 선택",
                                range(len(log_options)),
                                format_func=lambda x: log_options[x]
                            )
                            
                            if selected_log_idx is not None:
                                selected_log = quality_logs.data[selected_log_idx]
                                
                                st.markdown("---")
                                st.markdown("## 📋 상세 비교 분석")
                                
                                # 기본 정보
                                col1, col2, col3, col4 = st.columns(4)
                                col1.metric("회사명", selected_log.get('company_name', 'N/A'))
                                col2.metric("키워드 수", len(selected_log.get('selected_keywords', [])))
                                col3.metric("추출 성공률", f"{selected_log.get('extraction_success_rate', 0)}%")
                                col4.metric("표 인식", f"{selected_log.get('ocr_tables_count', 0)}개")
                                
                                st.markdown("---")
                                
                                # 3단계 비교 탭
                                comp_tab1, comp_tab2, comp_tab3, comp_tab4 = st.tabs([
                                    "1️⃣ 선택된 키워드", 
                                    "2️⃣ OCR 원본 데이터", 
                                    "3️⃣ LLM 추출 데이터",
                                    "4️⃣ 보고서 생성 결과"
                                ])
                                
                                with comp_tab1:
                                    st.markdown("### 📌 사용자가 선택한 추출 키워드")
                                    keywords = selected_log.get('selected_keywords', [])
                                    
                                    if keywords:
                                        cols = st.columns(4)
                                        for idx, kw in enumerate(keywords):
                                            col = cols[idx % 4]
                                            col.markdown(f"""
                                            <div style='
                                                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                                                color: white;
                                                padding: 10px;
                                                border-radius: 8px;
                                                text-align: center;
                                                margin: 5px 0;
                                                font-weight: 500;
                                            '>
                                                {kw}
                                            </div>
                                            """, unsafe_allow_html=True)
                                        st.caption(f"총 {len(keywords)}개의 키워드가 선택되었습니다")
                                    else:
                                        st.warning("키워드 정보가 없습니다")
                                
                                with comp_tab2:
                                    st.markdown("### 📄 OCR 원본 추출 데이터 (Upstage Parse)")
                                    
                                    # 표 데이터
                                    structured_data = selected_log.get('ocr_structured_data', {})
                                    if structured_data and structured_data.get('tables'):
                                        st.markdown(f"#### 📊 인식된 표 ({len(structured_data['tables'])}개)")
                                        for idx, table in enumerate(structured_data['tables']):
                                            with st.expander(f"표 {idx+1} (페이지 {table.get('page', '?')})"):
                                                st.text(table.get('content', '내용 없음')[:1000])
                                    else:
                                        st.info("구조화된 표 데이터가 없습니다")
                                    
                                    # 원본 텍스트
                                    st.markdown("#### 📝 추출된 원본 텍스트 (일부)")
                                    raw_text = selected_log.get('ocr_raw_text', '')
                                    if raw_text:
                                        st.text_area("OCR 원본", raw_text[:2000], height=300)
                                        st.caption(f"전체 길이: {len(raw_text)}자 (처음 2000자 표시)")
                                    else:
                                        st.warning("원본 텍스트가 없습니다")
                                
                                with comp_tab3:
                                    st.markdown("### 🤖 LLM이 추출한 데이터")
                                    
                                    extracted = selected_log.get('llm_extracted_data', {})
                                    if extracted:
                                        # 성공/실패 구분
                                        success_data = {k: v for k, v in extracted.items() if v and v != "정보 없음"}
                                        failed_data = {k: v for k, v in extracted.items() if not v or v == "정보 없음"}
                                        
                                        col1, col2 = st.columns(2)
                                        col1.metric("✅ 추출 성공", len(success_data))
                                        col2.metric("❌ 추출 실패", len(failed_data))
                                        
                                        st.markdown("---")
                                        
                                        # 성공한 데이터
                                        st.markdown("#### ✅ 성공적으로 추출된 데이터")
                                        if success_data:
                                            for key, value in success_data.items():
                                                st.markdown(f"**{key}**")
                                                st.markdown(f"""
                                                <div style='
                                                    background: #f0fdf4;
                                                    border-left: 4px solid #22c55e;
                                                    padding: 10px 15px;
                                                    margin: 5px 0 15px 0;
                                                    border-radius: 4px;
                                                '>
                                                    {value}
                                                </div>
                                                """, unsafe_allow_html=True)
                                        else:
                                            st.info("추출된 데이터가 없습니다")
                                        
                                        st.markdown("---")
                                        
                                        # 실패한 데이터
                                        if failed_data:
                                            st.markdown("#### ❌ 추출 실패 데이터")
                                            for key in failed_data.keys():
                                                st.markdown(f"""
                                                <div style='
                                                    background: #fef2f2;
                                                    border-left: 4px solid #ef4444;
                                                    padding: 10px 15px;
                                                    margin: 5px 0;
                                                    border-radius: 4px;
                                                    color: #991b1b;
                                                '>
                                                    <strong>{key}</strong>: 정보 없음
                                                </div>
                                                """, unsafe_allow_html=True)
                                        
                                        # LLM 메타데이터
                                        st.markdown("---")
                                        st.caption(f"모델: {selected_log.get('llm_model', 'N/A')} | "
                                                 f"처리 시간: {selected_log.get('llm_extraction_time_ms', 0)}ms")
                                    else:
                                        st.warning("추출된 데이터가 없습니다")
                                
                                with comp_tab4:
                                    st.markdown("### 📄 최종 생성된 보고서")
                                    
                                    if selected_log.get('report_generated'):
                                        report = selected_log.get('report_content', '')
                                        if report:
                                            st.markdown(report)
                                            
                                            st.markdown("---")
                                            st.caption(f"모델: {selected_log.get('report_model', 'N/A')} | "
                                                     f"생성 시간: {selected_log.get('report_generation_time_ms', 0)}ms | "
                                                     f"길이: {len(report)}자")
                                        else:
                                            st.warning("보고서 내용이 없습니다")
                                    else:
                                        st.info("아직 보고서가 생성되지 않았습니다")
                                
                                # 전체 비교 요약
                                st.markdown("---")
                                st.markdown("## 📊 종합 비교 요약")
                                
                                col1, col2, col3 = st.columns(3)
                                
                                with col1:
                                    st.markdown("### 1️⃣ OCR 단계")
                                    st.metric("표 인식", f"{selected_log.get('ocr_tables_count', 0)}개")
                                    st.metric("차트/그래프 인식", f"{selected_log.get('ocr_charts_count', 0)}개")
                                    st.metric("텍스트 길이", f"{len(selected_log.get('ocr_raw_text', ''))}자")
                                
                                with col2:
                                    st.markdown("### 2️⃣ LLM 추출")
                                    st.metric("성공", selected_log.get('keywords_with_data', 0))
                                    st.metric("실패", selected_log.get('keywords_missing_data', 0))
                                    st.metric("성공률", f"{selected_log.get('extraction_success_rate', 0)}%")
                                
                                with col3:
                                    st.markdown("### 3️⃣ 보고서 생성")
                                    if selected_log.get('report_generated'):
                                        st.success("✅ 생성 완료")
                                        st.metric("길이", f"{len(selected_log.get('report_content', ''))}자")
                                    else:
                                        st.error("❌ 미생성")
                                
                                # TXT 파일로 내보내기 버튼
                                st.markdown("---")
                                st.markdown("## 📥 AI 분석용 TXT 파일 내보내기")
                                st.info("💡 이 로그를 TXT 파일로 다운로드하여 AI에게 첨부하면, 자동으로 문제점을 분석하고 개선 방안을 제시받을 수 있습니다.")
                                
                                if st.button("📥 TXT 파일로 내보내기", type="primary", use_container_width=True):
                                    # TXT 파일 생성
                                    txt_content = generate_quality_log_txt(selected_log)
                                    
                                    # 파일명 생성
                                    company = selected_log.get('company_name', 'Unknown').replace(' ', '_')
                                    created_at = selected_log.get('created_at', '')[:10]
                                    filename = f"quality_log_{company}_{created_at}.txt"
                                    
                                    # 다운로드 버튼
                                    st.download_button(
                                        label="💾 다운로드",
                                        data=txt_content,
                                        file_name=filename,
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                    
                                    st.success("✅ TXT 파일이 생성되었습니다! 다운로드 버튼을 클릭하세요.")
                                    
                                    with st.expander("📋 파일 미리보기"):
                                        st.text(txt_content[:2000] + "\n\n... (전체 내용은 다운로드하세요)")
                    
                    except Exception as e:
                        st.error(f"데이터 품질 로그 조회 실패: {e}")
                        st.code(traceback.format_exc())